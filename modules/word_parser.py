"""
modules/word_parser.py

DOCX parsing: document-type classification, block extraction, field extraction.
Drop into modules/ — app2.py already imports parse_word() from here.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx2txt
from docx import Document


# ─────────────────────────────────────────────────────────────────────────────
# Document Classification
# ─────────────────────────────────────────────────────────────────────────────

_DOC_TYPE_RULES: list[tuple[str, list[str]]] = [
    ("Insurance Policy",          ["policy number", "insured", "premium", "coverage", "deductible", "endorsement"]),
    ("Claims Report",             ["claim number", "claimant", "date of loss", "loss date", "claim status", "adjuster"]),
    ("Loss Run Report",           ["loss run", "total incurred", "reserve", "total paid", "open claims", "closed claims"]),
    ("Medical Report",            ["diagnosis", "patient", "physician", "treatment", "medication", "icd"]),
    ("Legal Document",            ["plaintiff", "defendant", "court", "jurisdiction", "attorney", "whereas"]),
    ("Invoice / Bill",            ["invoice number", "amount due", "payment terms", "bill to", "subtotal"]),
    ("Certificate of Insurance",  ["certificate", "certificate holder", "additional insured", "acord"]),
    ("Correspondence / Letter",   ["dear", "sincerely", "regards", "to whom it may concern"]),
    ("Contract / Agreement",      ["agreement", "terms and conditions", "parties agree", "effective date", "termination"]),
    ("Financial Statement",       ["balance sheet", "income statement", "cash flow", "assets", "liabilities"]),
]

_DOC_TYPE_FALLBACK = "General Document"


def classify_word_document(raw_text: str) -> dict:
    """
    Classify the Word document by text-signal matching.

    Returns:
        {
            "doc_type_label": "Claims Report",
            "confidence": 0.83,
            "matched_signals": ["claim number", "date of loss"]
        }
    """
    lower = raw_text.lower()
    best_label   = _DOC_TYPE_FALLBACK
    best_score   = 0.0
    best_signals: list[str] = []

    for label, keywords in _DOC_TYPE_RULES:
        hits  = [kw for kw in keywords if kw in lower]
        score = len(hits) / len(keywords)
        if score > best_score:
            best_score   = score
            best_label   = label
            best_signals = hits

    return {
        "doc_type_label":  best_label,
        "confidence":      round(best_score, 2),
        "matched_signals": best_signals,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Public API  (app2.py calls parse_word)
# ─────────────────────────────────────────────────────────────────────────────

def parse_word(file_path: str | Path, llm_client=None) -> dict:
    """
    Parse a .docx and return a structure compatible with the rest of the pipeline.

    Returns:
        {
            "doc_type":          "word_document",
            "doc_label":         "Word Document",
            "doc_classification": { doc_type_label, confidence, matched_signals },
            "raw_text":          "...",
            "blocks":            [...],
            "fields":            [...],
        }
    """
    blocks   = extract_word_blocks(file_path)
    raw_text = "\n".join(b["text"] for b in blocks if b.get("text", "").strip()).strip()

    return {
        "doc_type":           "word_document",
        "doc_label":          "Word Document",
        "doc_classification": classify_word_document(raw_text),
        "raw_text":           raw_text,
        "blocks":             blocks,
        "fields":             extract_word_fields_from_blocks(blocks, llm_client=llm_client),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Helpers used by app2.py  (get_sheet_names / get_sheet_dimensions equivalents)
# ─────────────────────────────────────────────────────────────────────────────

def get_word_sheet_names(file_path: str | Path) -> list[str]:
    return ["Document"]


def get_word_dimensions(file_path: str | Path, section_name: str = "Document") -> tuple[int, int]:
    blocks = extract_word_blocks(file_path)
    return len(blocks), 1


def extract_word_text(file_path: str | Path) -> str:
    try:
        return (docx2txt.process(str(file_path)) or "").strip()
    except Exception:
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Block extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_word_blocks(file_path: str | Path) -> list[dict]:
    """
    Return ordered blocks (paragraphs + table rows/cells) for traceability.

    Each block:
        block_id, block_type, text,
        para_index, table_index, row_index, col_index
    """
    doc    = Document(str(file_path))
    blocks: list[dict] = []
    bid    = 1

    # Paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        txt = (para.text or "").strip()
        if txt:
            blocks.append({
                "block_id":    bid,   "block_type":  "paragraph",
                "text":        txt,   "para_index":  p_idx,
                "table_index": None,  "row_index":   None,  "col_index": None,
            })
            bid += 1

    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells    = [(c.text or "").strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                blocks.append({
                    "block_id":    bid,   "block_type":  "table_row",
                    "text":        row_text, "para_index": None,
                    "table_index": t_idx, "row_index":   r_idx, "col_index": None,
                })
                bid += 1

            for c_idx, cell in enumerate(row.cells):
                txt = (cell.text or "").strip()
                if txt:
                    blocks.append({
                        "block_id":    bid,   "block_type":  "table_cell",
                        "text":        txt,   "para_index":  None,
                        "table_index": t_idx, "row_index":   r_idx, "col_index": c_idx,
                    })
                    bid += 1

    return blocks


# ─────────────────────────────────────────────────────────────────────────────
# Field extraction
# ─────────────────────────────────────────────────────────────────────────────

_LABEL_KEYWORDS = [
    "policy", "claim", "insured", "carrier", "date", "status",
    "premium", "limit", "deductible", "period", "name", "number",
    "less", "add", "tax",
]

_CANONICAL_MAP: dict[str, str] = {
    "claim number": "Claim Number",  "claim no":   "Claim Number",
    "claim #":      "Claim Number",  "claim id":   "Claim Number",
    "policy number":"Policy Number", "policy no":  "Policy Number",
    "policy #":     "Policy Number", "insured":    "Insured",
    "carrier":      "Carrier",       "loss date":  "Loss Date",
    "date of loss": "Loss Date",     "date reported": "Date Reported",
    "status":       "Status",        "claimant":   "Claimant Name",
    "claimant name":"Claimant Name", "description of loss": "Description of Loss",
    "total paid":   "Total Paid",    "reserve":    "Reserve",
    "total incurred":"Total Incurred","effective date": "Effective Date",
    "expiration date":"Expiration Date","lob":      "Line of Business",
    "line of business":"Line of Business",
}

_LABEL_PATTERNS = set(_CANONICAL_MAP.keys())


def _canon(label: str) -> str:
    k = re.sub(r"\s+", " ", label.strip().lower())
    return _CANONICAL_MAP.get(k, label.strip().title())


def _is_label(text: str) -> bool:
    if not text:
        return False
    t = text.strip().lower()
    return len(t) <= 60 and any(kw in t for kw in _LABEL_KEYWORDS)


def _make_field(name: str, value: str, conf: float, block: dict) -> dict:
    return {
        "field_name":   name,
        "value":        value,
        "confidence":   conf,
        "source_block": block.get("block_id"),
        "source_text":  block.get("text", ""),
        "source_para":  block.get("para_index"),
        "source_table": block.get("table_index"),
        "source_row":   block.get("row_index"),
        "source_col":   block.get("col_index"),
    }


def extract_word_fields_from_blocks(blocks: list[dict], llm_client=None) -> list[dict]:
    """
    Four-pass field extractor:
      1. Inline  "Label: Value"  text patterns
      2. Table rows  "Label | Value"
      3. Adjacent table-cell pairs  (col N → col N+1)
      4. Fallback regex across full text
    """
    fields: list[dict] = []
    seen:   set[str]   = set()

    # ── Pass 1: "Label: Value" ────────────────────────────────────────────
    _pat = re.compile(r"^\s*([A-Za-z0-9 #/_\-\(\)\.]{2,60})\s*:\s*(.+?)\s*$")
    for b in blocks:
        m = _pat.match(b.get("text", ""))
        if not m:
            continue
        fname = _canon(m.group(1))
        val   = m.group(2).strip()
        if val and fname not in seen:
            fields.append(_make_field(fname, val, 0.95, b))
            seen.add(fname)

    # ── Pass 2: table_row "Label | Value" ────────────────────────────────
    for b in blocks:
        if b.get("block_type") != "table_row":
            continue
        parts = [p.strip() for p in b["text"].split("|") if p.strip()]
        if len(parts) >= 2 and parts[0].lower() in _LABEL_PATTERNS:
            fname = _canon(parts[0])
            if fname not in seen:
                fields.append(_make_field(fname, parts[1], 0.92, b))
                seen.add(fname)

    # ── Pass 3: adjacent table_cell pairs ────────────────────────────────
    cells = [b for b in blocks if b.get("block_type") == "table_cell"]
    for i in range(len(cells) - 1):
        L, R = cells[i], cells[i + 1]
        same_row = (
            L.get("table_index") == R.get("table_index")
            and L.get("row_index") == R.get("row_index")
            and R.get("col_index") == (L.get("col_index") or 0) + 1
        )
        if not same_row:
            continue
        raw_lbl = (L.get("text") or "").strip()
        raw_val = (R.get("text") or "").strip()
        if _is_label(raw_lbl) and raw_val and len(raw_val) <= 200:
            fname = _canon(raw_lbl)
            if fname not in seen:
                merged_block = {**R, "text": f"{raw_lbl}: {raw_val}"}
                fields.append(_make_field(fname, raw_val, 0.93, merged_block))
                seen.add(fname)

    # ── Pass 4: fallback regex ────────────────────────────────────────────
    if not fields:
        joined = "\n".join(b.get("text", "") for b in blocks)
        for pat, name in [
            (r"\bPolicy\s*(?:Number|No|#)\s*[:\-]?\s*([A-Z0-9\-_\/]+)", "Policy Number"),
            (r"\bClaim\s*(?:Number|No|#|ID)\s*[:\-]?\s*([A-Z0-9\-_\/]+)", "Claim Number"),
            (r"\bInsured\s*[:\-]?\s*([^\n\r]+)",                           "Insured"),
            (r"\bCarrier\s*[:\-]?\s*([^\n\r]+)",                           "Carrier"),
            (r"\bLoss\s*Date\s*[:\-]?\s*([^\n\r]+)",                       "Loss Date"),
            (r"\bEffective\s*Date\s*[:\-]?\s*([^\n\r]+)",                  "Effective Date"),
        ]:
            m = re.search(pat, joined, re.IGNORECASE)
            if m:
                val = m.group(1).strip()
                fields.append({
                    "field_name": name, "value": val, "confidence": 0.80,
                    "source_block": None, "source_text": val,
                    "source_para": None, "source_table": None,
                    "source_row": None,  "source_col":  None,
                })

    return fields