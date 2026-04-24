"""
ui/word_analysis.py  — v8

CHANGES vs v7:
──────────────────────────────────────────────────────────────────────────────
1. AI badge removed — the purple "AI" tag is no longer shown on field names.
2. Signals auto-run — signals are detected automatically when the Signals tab
   is first opened; no manual "Run" button click needed. A "Re-run Analysis"
   button is still shown for subsequent reruns.
3. Legal-document field extraction enhanced — Case Summary and Causes of Action
   (multi-line) are now extracted as proper fields via both the direct docx
   parser and the LLM field-extraction prompt.
4. AttributeError fix — incorrect_fields items in the AI Assistant dimension
   renderer now safely handle both dict and plain-string entries.
──────────────────────────────────────────────────────────────────────────────
"""

from __future__ import annotations

import datetime
import html
import json
import os
import re

import streamlit as st
import streamlit.components.v1 as components


# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

_DOC_TYPE_META = {
    "Insurance Policy":          {"icon": "🛡️",  "color": "#0369a1", "bg": "rgba(3,105,161,0.06)"},
    "Claims Report":             {"icon": "🚨",  "color": "#dc2626", "bg": "rgba(220,38,38,0.06)"},
    "Loss Run Report":           {"icon": "📊",  "color": "#059669", "bg": "rgba(5,150,105,0.06)"},
    "Medical Report":            {"icon": "🏥",  "color": "#2563eb", "bg": "rgba(37,99,235,0.06)"},
    "Legal Document":            {"icon": "⚖️",  "color": "#7c3aed", "bg": "rgba(124,58,237,0.06)"},
    "Invoice / Bill":            {"icon": "🧾",  "color": "#b45309", "bg": "rgba(180,83,9,0.06)"},
    "Certificate of Insurance":  {"icon": "📜",  "color": "#0891b2", "bg": "rgba(8,145,178,0.06)"},
    "Correspondence / Letter":   {"icon": "✉️",  "color": "#6b7280", "bg": "rgba(107,114,128,0.06)"},
    "Contract / Agreement":      {"icon": "📋",  "color": "#7c3aed", "bg": "rgba(124,58,237,0.06)"},
    "Financial Statement":       {"icon": "💰",  "color": "#059669", "bg": "rgba(5,150,105,0.06)"},
    "General Document":          {"icon": "📄",  "color": "#475569", "bg": "rgba(71,85,105,0.06)"},
}

# ── Four-tier taxonomy ────────────────────────────────────────────────────────
_TAXONOMY = {
    "Highly Severe": {"color": "#dc2626", "bg": "rgba(220,38,38,0.06)",  "icon": "🔥"},
    "High":          {"color": "#ea580c", "bg": "rgba(234,88,12,0.06)",  "icon": "🔴"},
    "Moderate":      {"color": "#ca8a04", "bg": "rgba(202,138,4,0.06)",  "icon": "🟡"},
    "Low":           {"color": "#16a34a", "bg": "rgba(22,163,74,0.06)",  "icon": "🟢"},
}

_BG      = "#ffffff"
_BG2     = "#f8f9fa"
_BORDER  = "#e2e8f0"
_BORDER2 = "#cbd5e1"
_TXT     = "#0f172a"
_TXT2    = "#1e293b"
_LBL     = "#64748b"
_LBL2    = "#94a3b8"


# ─────────────────────────────────────────────────────────────────────────────
# AZURE OPENAI CLIENT FACTORY
# ─────────────────────────────────────────────────────────────────────────────

def _get_oai_client():
    try:
        from openai import AzureOpenAI
        return AzureOpenAI(
            azure_endpoint=os.environ.get("OPENAI_DEPLOYMENT_ENDPOINT", ""),
            api_key=os.environ.get("OPENAI_API_KEY", ""),
            api_version=os.environ.get("OPENAI_API_VERSION", "2024-12-01-preview"),
        )
    except Exception:
        return None


def _std_model() -> str:
    return os.environ.get("OPENAI_DEPLOYMENT_NAME", "gpt-4o-mini")


def _adv_model() -> str:
    return os.environ.get("OPENAI_DEPLOYMENT_NAME_ADV", "gpt-4o")


def _chat(client, model: str, system: str, user: str,
          max_tokens: int = 1200, temperature: float = 0.0) -> str | None:
    if client is None:
        return None
    try:
        resp = client.chat.completions.create(
            model=model,
            max_tokens=max_tokens,
            temperature=temperature,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user},
            ],
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception:
        return None


def _parse_json_response(raw: str | None) -> dict | list | None:
    if not raw:
        return None
    clean = re.sub(r"^```(?:json)?\s*|```\s*$", "", raw, flags=re.MULTILINE).strip()
    try:
        return json.loads(clean)
    except Exception:
        return None


# ─────────────────────────────────────────────────────────────────────────────
# DIRECT FIELD EXTRACTION FROM DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _is_paired_table(rows: list) -> bool:
    non_empty = [
        [c.text.strip() for c in r.cells]
        for r in rows
        if any(c.text.strip() for c in r.cells)
    ]
    if len(non_empty) < 2:
        return False
    upper_count = sum(
        1 for cells in non_empty
        if all(c == c.upper() for c in cells if c)
    )
    if upper_count / len(non_empty) < 0.40:
        return False
    pairs = 0
    for idx in range(len(non_empty) - 1):
        a_upper = all(c == c.upper() for c in non_empty[idx]     if c)
        b_upper = all(c == c.upper() for c in non_empty[idx + 1] if c)
        if a_upper and not b_upper:
            pairs += 1
    return pairs >= 2


def _extract_fields_from_docx(file_path: str) -> list[dict]:
    """
    Extract key-value fields from tables and paragraphs in the docx.
    Also extracts multi-line fields like Case Complaint Summary and
    Causes of Action that appear as labelled sections in the document body.
    """
    try:
        from docx import Document
    except ImportError:
        return []
    try:
        doc = Document(file_path)
    except Exception:
        return []

    fields: list[dict] = []
    seen_keys: set[str] = set()
    _SKIP_NAMES = {"field", "details", "key", "value", "source", "trellis"}

    def _add(name: str, value: str, conf: float = 0.9, **meta) -> None:
        name  = name.strip().strip(":").strip()
        value = value.strip()
        if not name or not value:
            return
        if name.lower() in _SKIP_NAMES:
            return
        key_norm = name.upper()
        if key_norm in seen_keys:
            return
        seen_keys.add(key_norm)
        fields.append({
            "field_name":   name,
            "value":        value,
            "confidence":   conf,
            "source_text":  value,
            "source_block": None,
            "source_para":  meta.get("source_para"),
            "source_table": meta.get("source_table"),
            "source_row":   meta.get("source_row"),
            "source_col":   meta.get("source_col"),
        })

    # ── Table extraction ──────────────────────────────────────────────────────
    for t_idx in range(len(doc.tables) - 1, -1, -1):
        table  = doc.tables[t_idx]
        r_list = list(table.rows)
        if _is_paired_table(r_list):
            non_empty = [
                (i, [c.text.strip() for c in r.cells])
                for i, r in enumerate(r_list)
                if any(c.text.strip() for c in r.cells)
            ]
            idx = 0
            while idx < len(non_empty):
                i, cells_a = non_empty[idx]
                a_upper = all(c == c.upper() for c in cells_a if c)
                if a_upper and idx + 1 < len(non_empty):
                    j, cells_b = non_empty[idx + 1]
                    b_upper = all(c == c.upper() for c in cells_b if c)
                    if not b_upper:
                        for col, (hdr, val) in enumerate(zip(cells_a, cells_b)):
                            if hdr and val:
                                _add(hdr, val, conf=0.93,
                                     source_table=t_idx, source_row=j, source_col=col)
                        idx += 2
                        continue
                idx += 1
        else:
            for r_idx, row in enumerate(r_list):
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                key_cell, value_cell = cells[0], cells[1]
                if not key_cell or not value_cell:
                    continue
                if key_cell == value_cell:
                    continue
                _add(key_cell, value_cell, conf=0.95,
                     source_table=t_idx, source_row=r_idx, source_col=1)

    # ── Paragraph extraction — inline key: value pairs ────────────────────────
    paragraphs = doc.paragraphs

    # Section labels that introduce multi-line content blocks
    _SECTION_LABELS = {
        "cause(s) of action":       "Causes of Action",
        "causes of action":         "Causes of Action",
        "case complaint summary":   "Case Complaint Summary",
        "complaint summary":        "Case Complaint Summary",
        "overview":                 None,   # skip — too generic
        "case details":             None,
    }

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if not text or text.startswith("http") or len(text) > 300:
            i += 1
            continue

        # Check for known multi-line section headers
        text_lower = text.lower().rstrip(":")
        if text_lower in _SECTION_LABELS:
            mapped_name = _SECTION_LABELS[text_lower]
            if mapped_name:
                # Collect subsequent non-empty, non-header paragraphs as the value
                value_lines = []
                j = i + 1
                while j < len(paragraphs):
                    next_text = paragraphs[j].text.strip()
                    # Stop at next section header or empty gap after content
                    if next_text.lower().rstrip(":") in _SECTION_LABELS:
                        break
                    if next_text.startswith("Source:"):
                        break
                    if next_text:
                        value_lines.append(next_text)
                    j += 1
                if value_lines:
                    _add(mapped_name, " | ".join(value_lines), conf=0.90,
                         source_para=i, source_table=None, source_row=None, source_col=None)
            i += 1
            continue

        # Standard inline key: value pair
        if ": " in text:
            parts = text.split(": ", 1)
            if len(parts) == 2 and parts[0] and parts[1] and len(parts[0]) < 60:
                _add(parts[0], parts[1], conf=0.85,
                     source_para=i, source_table=None, source_row=None, source_col=None)

        i += 1

    return fields


# ─────────────────────────────────────────────────────────────────────────────
# LLM-POWERED FIELD EXTRACTION  (document-type-aware)
# ─────────────────────────────────────────────────────────────────────────────

_DOC_TYPE_FIELDS = {
    "Insurance Policy": [
        "Policy Number", "Policy Holder", "Insured Name", "Carrier", "Underwriter",
        "Effective Date", "Expiration Date", "Premium Amount", "Coverage Type",
        "Deductible", "Policy Limit", "Named Insured", "Agent", "Endorsements",
    ],
    "Claims Report": [
        "Claim Number", "Claimant Name", "Policy Number", "Carrier", "Loss Date",
        "Report Date", "Adjuster", "Claim Status", "Incident Description",
        "Total Incurred", "Reserve Amount", "Paid Amount", "Coverage",
    ],
    "Loss Run Report": [
        "Policy Number", "Insured", "Carrier", "Policy Period", "Total Claims",
        "Total Paid", "Total Reserves", "Total Incurred", "Number of Occurrences",
        "Largest Claim", "Report Date", "Experience Period",
    ],
    "Medical Report": [
        "Patient Name", "Date of Birth", "Patient ID", "Physician", "Diagnosis",
        "ICD Code", "Date of Service", "Facility", "Treatment", "Medications",
        "Follow-up Date", "Insurance ID", "Referring Physician",
    ],
    "Legal Document": [
        "Parties", "Plaintiff", "Defendant", "Case Number", "Court", "Judge",
        "Filing Date", "Jurisdiction", "Subject Matter", "Effective Date",
        "Governing Law", "Arbitration Clause", "Signatures", "Notary",
        "Case Category", "Practice Area", "Matter Type", "Case Status",
        "Policy Number", "Causes of Action", "Case Complaint Summary",
    ],
    "Invoice / Bill": [
        "Invoice Number", "Invoice Date", "Due Date", "Vendor", "Client",
        "Total Amount", "Tax Amount", "Line Items", "Payment Terms",
        "Account Number", "Purchase Order",
    ],
    "Certificate of Insurance": [
        "Certificate Number", "Insured", "Carrier", "Policy Number",
        "Effective Date", "Expiration Date", "Coverage Type", "Policy Limit",
        "Certificate Holder", "Producer",
    ],
    "Financial Statement": [
        "Entity Name", "Period", "Total Revenue", "Total Expenses", "Net Income",
        "Total Assets", "Total Liabilities", "Equity", "Auditor", "Fiscal Year",
    ],
    "Contract / Agreement": [
        "Parties", "Effective Date", "Expiration Date", "Contract Value",
        "Payment Terms", "Governing Law", "Termination Clause", "Signatures",
    ],
}


def _extract_fields_llm(doc_label: str, raw_text: str) -> list[dict]:
    if not raw_text.strip():
        return []

    target_fields = _DOC_TYPE_FIELDS.get(doc_label, [])
    fields_hint = (
        f"Focus especially on these fields if present: {', '.join(target_fields)}. "
        "For 'Causes of Action', concatenate all causes into a single string separated by ' | '. "
        "For 'Case Complaint Summary', provide the full summary text."
        if target_fields else
        "Extract all meaningful key-value pairs you can identify."
    )

    system = (
        "You are a precise document data-extraction engine. "
        "Given document text, extract key-value pairs relevant to the document type. "
        "Return ONLY a JSON array — no preamble, no markdown fences — like:\n"
        '[{"field_name":"Policy Number","value":"POL-2024-001","confidence":0.95},'
        ' {"field_name":"Insured","value":"Acme Corp","confidence":0.90}]'
    )
    user = (
        f"Document type: {doc_label}\n\n"
        f"{fields_hint}\n\n"
        f"Document text (may be truncated):\n{raw_text[:5000]}"
        + ("\n[...truncated...]" if len(raw_text) > 5000 else "")
    )

    client = _get_oai_client()
    raw = _chat(client, _std_model(), system, user, max_tokens=1500)
    parsed = _parse_json_response(raw)

    if not isinstance(parsed, list):
        return []

    result = []
    for item in parsed:
        fn  = str(item.get("field_name", "")).strip()
        val = str(item.get("value", "")).strip()
        conf = float(item.get("confidence", 0.85))
        if fn and val:
            result.append({
                "field_name":   fn,
                "value":        val,
                "confidence":   min(max(conf, 0.0), 1.0),
                "source_text":  val,
                "source_block": None,
                "source_para":  None,
                "source_table": None,
                "source_row":   None,
                "source_col":   None,
            })
    return result


def _merge_fields(existing: list[dict], direct: list[dict]) -> list[dict]:
    direct_keys = {f["field_name"].upper() for f in direct}
    extra = [f for f in existing if f.get("field_name", "").upper() not in direct_keys]
    return direct + extra


# ─────────────────────────────────────────────────────────────────────────────
# STYLE HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _conf_badge(conf: float) -> str:
    pct = int(conf * 100)
    c   = "#16a34a" if pct >= 80 else "#ca8a04" if pct >= 60 else "#dc2626"
    bg  = "#f0fdf4" if pct >= 80 else "#fefce8" if pct >= 60 else "#fef2f2"
    return (
        f"<span style='background:{bg};border:1px solid {c}40;border-radius:20px;"
        f"padding:1px 8px;font-size:10px;color:{c};font-weight:700;"
        f"font-family:monospace;'>{pct}%</span>"
    )


def _score_badge(score: int) -> str:
    c  = "#16a34a" if score >= 80 else "#ca8a04" if score >= 60 else "#dc2626"
    bg = "#f0fdf4" if score >= 80 else "#fefce8" if score >= 60 else "#fef2f2"
    return (
        f"<span style='background:{bg};border:1px solid {c}40;border-radius:20px;"
        f"padding:2px 10px;font-size:11px;color:{c};font-weight:700;"
        f"font-family:monospace;'>{score}/100</span>"
    )


def _verdict_badge(verdict: str) -> str:
    vmap = {
        "Pass":          ("#16a34a", "#f0fdf4"),
        "Validated":     ("#16a34a", "#f0fdf4"),
        "Adequate":      ("#16a34a", "#f0fdf4"),
        "Credible":      ("#16a34a", "#f0fdf4"),
        "Fail":          ("#dc2626", "#fef2f2"),
        "Failed":        ("#dc2626", "#fef2f2"),
        "Unsupported":   ("#dc2626", "#fef2f2"),
        "Critical Gaps": ("#dc2626", "#fef2f2"),
        "Review":        ("#ca8a04", "#fefce8"),
        "Needs Review":  ("#ca8a04", "#fefce8"),
        "Questionable":  ("#ca8a04", "#fefce8"),
        "Gaps Identified": ("#ca8a04", "#fefce8"),
    }
    c, bg = vmap.get(verdict, ("#64748b", "#f8fafc"))
    return (
        f"<span style='background:{bg};border:1px solid {c}40;border-radius:6px;"
        f"padding:3px 12px;font-size:11px;color:{c};font-weight:700;"
        f"font-family:monospace;'>{verdict}</span>"
    )


def _section_header(title: str, subtitle: str = "") -> str:
    sub = (
        f"<span style='font-size:10px;color:{_LBL};font-family:monospace;'>{subtitle}</span>"
        if subtitle else ""
    )
    return (
        f"<div style='display:flex;align-items:center;gap:10px;margin-bottom:14px;'>"
        f"<div style='font-size:11px;font-weight:700;color:{_TXT2};font-family:monospace;"
        f"text-transform:uppercase;letter-spacing:1.5px;white-space:nowrap;'>{title}</div>"
        f"{sub}"
        f"<div style='flex:1;height:1px;background:linear-gradient(90deg,{_BORDER},{_BG});'>"
        f"</div></div>"
    )


def _card(content: str, border_color: str = _BORDER, bg: str = _BG) -> str:
    return (
        f"<div style='background:{bg};border:1px solid {border_color};"
        f"border-radius:8px;padding:14px 16px;margin-bottom:10px;'>{content}</div>"
    )


# ─────────────────────────────────────────────────────────────────────────────
# SESSION-STATE HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _wa_key(uploaded_name: str, suffix: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_]", "_", uploaded_name)
    return f"_wa_{safe}_{suffix}"


def _edits(uploaded_name: str) -> dict:
    k = _wa_key(uploaded_name, "edits")
    if k not in st.session_state:
        st.session_state[k] = {}
    return st.session_state[k]


def _edit_history(uploaded_name: str) -> dict:
    k = _wa_key(uploaded_name, "edit_hist")
    if k not in st.session_state:
        st.session_state[k] = {}
    return st.session_state[k]


def _sync_edit(field_name: str, new_value: str, uploaded_name: str) -> None:
    eds  = _edits(uploaded_name)
    hist = _edit_history(uploaded_name)
    old  = eds.get(field_name, "")
    eds[field_name] = new_value
    if field_name not in hist:
        hist[field_name] = []
    if not hist[field_name] or hist[field_name][-1]["to"] != new_value:
        hist[field_name].append({
            "timestamp": datetime.datetime.now().isoformat(),
            "from": old, "to": new_value,
        })


# ─────────────────────────────────────────────────────────────────────────────
# DOCX → HTML  (pure python-docx, NO mammoth)
# ─────────────────────────────────────────────────────────────────────────────

def _para_to_html(para, bi: int) -> str:
    style_name = (para.style.name or "").lower()
    if "heading 1" in style_name:
        tag = "h1"
    elif "heading 2" in style_name:
        tag = "h2"
    elif "heading 3" in style_name or "heading 4" in style_name:
        tag = "h3"
    else:
        tag = "p"

    inner_parts: list[str] = []
    for run in para.runs:
        txt = html.escape(run.text)
        if not txt:
            continue
        if run.bold:
            txt = f"<strong>{txt}</strong>"
        if run.italic:
            txt = f"<em>{txt}</em>"
        if run.underline:
            txt = f"<u>{txt}</u>"
        inner_parts.append(txt)

    inner = "".join(inner_parts) or "&nbsp;"
    return f'<{tag} data-bi="{bi}">{inner}</{tag}>'


def _table_to_html(table, counter: list) -> str:
    rows_html: list[str] = []
    for row in table.rows:
        bi = counter[0]
        counter[0] += 1
        cells_html = "".join(
            f"<td>{html.escape(cell.text)}</td>"
            for cell in row.cells
        )
        rows_html.append(f'<tr data-bi="{bi}">{cells_html}</tr>')
    return "<table>" + "".join(rows_html) + "</table>"


def _docx_to_html_indexed(file_path: str) -> tuple[str, int]:
    try:
        from docx import Document
    except ImportError:
        return "<p data-bi='0'>python-docx not installed.</p>", 1
    try:
        doc = Document(file_path)
    except Exception as exc:
        return f"<p data-bi='0'>Could not open document: {html.escape(str(exc))}</p>", 1

    counter = [0]
    parts:  list[str] = []
    para_map  = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}
    body = doc.element.body

    for child in body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p":
            para = para_map.get(id(child))
            if para is None:
                raw_txt = html.escape("".join(t.text or "" for t in child.iter()
                                              if t.tag.endswith("}t")))
                parts.append(f'<p data-bi="{counter[0]}">{raw_txt or "&nbsp;"}</p>')
            else:
                parts.append(_para_to_html(para, counter[0]))
            counter[0] += 1
        elif local == "tbl":
            tbl_obj = table_map.get(id(child))
            if tbl_obj is None:
                counter[0] += 1
                continue
            parts.append(_table_to_html(tbl_obj, counter))
        elif local == "sectPr":
            pass
        else:
            parts.append(f'<p data-bi="{counter[0]}">&nbsp;</p>')
            counter[0] += 1

    return "\n".join(parts), counter[0]


# ─────────────────────────────────────────────────────────────────────────────
# HIGHLIGHT TEXT RESOLUTION
# ─────────────────────────────────────────────────────────────────────────────

def _resolve_highlight_text(field: dict, current_val: str) -> str:
    val = (field.get("value") or current_val or field.get("source_text") or "").strip()
    return val


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT VIEWER
# ─────────────────────────────────────────────────────────────────────────────

def _render_word_document_viewer(
    file_path: str,
    field: dict,
    current_val: str,
    uploaded_name: str,
    viewer_height: int = 540,
) -> None:
    conf     = float(field.get("confidence", 0.9))
    conf_pct = int(conf * 100)
    fname    = field.get("field_name", "")

    if conf_pct >= 80:
        box_c = "#16a34a"; bdg_bg = "#f0fdf4"; bdg_fg = "#15803d"
    elif conf_pct >= 60:
        box_c = "#ca8a04"; bdg_bg = "#fefce8"; bdg_fg = "#92400e"
    else:
        box_c = "#dc2626"; bdg_bg = "#fef2f2"; bdg_fg = "#991b1b"

    indexed_html, _ = _docx_to_html_indexed(file_path)
    highlight_text  = _resolve_highlight_text(field, current_val)
    hl_js = highlight_text.replace("\\", "\\\\").replace('"', '\\"').replace("\n", " ")
    fn_js = fname.replace("\\", "\\\\").replace('"', '\\"')

    iframe = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
html, body {{ height: 100%; background: #dde1e7; font-family: 'Segoe UI', Arial, sans-serif; overflow-x: hidden; }}
#toolbar {{
  position: sticky; top: 0; z-index: 200;
  background: #1e293b; color: #f1f5f9;
  padding: 7px 16px;
  display: flex; align-items: center; gap: 10px;
  font-size: 11px; font-family: monospace;
  border-bottom: 2px solid {box_c};
}}
#toolbar .dn  {{ font-weight:700; color:#f8fafc; max-width:220px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }}
#toolbar .fc  {{ background:{box_c}22; border:1px solid {box_c}80; border-radius:20px; padding:2px 9px; color:{box_c}; font-weight:700; white-space:nowrap; }}
#toolbar .cc  {{ background:{bdg_bg}; border:1px solid {box_c}60; border-radius:20px; padding:2px 9px; color:{bdg_fg}; font-weight:800; white-space:nowrap; }}
#toolbar .sp  {{ flex:1; }}
#toolbar .ll  {{ font-size:10px; color:#94a3b8; }}
#wrap {{ max-width:800px; margin:16px auto 40px; padding:0 14px; }}
#page {{
  background:#fff;
  box-shadow: 0 2px 20px rgba(0,0,0,0.18);
  border-radius:3px;
  padding:52px 68px 68px;
  position:relative;
  min-height:500px;
}}
#page p  {{ font-size:13px; line-height:1.8; color:#1e293b; margin-bottom:5px; padding:2px 3px; border-radius:3px; }}
#page h1 {{ font-size:20px; font-weight:800; color:#0f172a; margin:16px 0 8px; padding:2px 3px; border-radius:3px; }}
#page h2 {{ font-size:16px; font-weight:700; color:#0f172a; margin:14px 0 6px; padding:2px 3px; border-radius:3px; }}
#page h3 {{ font-size:14px; font-weight:700; color:#0f172a; margin:12px 0 5px; padding:2px 3px; border-radius:3px; }}
#page table {{ width:100%; border-collapse:collapse; margin:10px 0; font-size:13px; }}
#page td, #page th {{ border:1px solid #cbd5e1; padding:6px 10px; color:#1e293b; vertical-align:top; }}
#page tr:nth-child(even) td {{ background:#f8fafc; }}
#page strong {{ color:#0f172a; }}
#page em     {{ font-style:italic; }}
#page u      {{ text-decoration:underline; }}
#page ul, #page ol {{ padding-left:20px; margin-bottom:8px; }}
#page li {{ font-size:13px; line-height:1.7; color:#1e293b; }}
#page p, #page h1, #page h2, #page h3, #page tr {{ opacity: 0.42; transition: opacity 0.2s; }}
.hlm {{ background:#fde047; color:#111827; border-radius:3px; padding:0 2px; font-weight:700; }}
.bbox-box {{
  position:absolute; pointer-events:none;
  border:2.5px solid {box_c};
  border-radius:6px;
  box-shadow: 0 0 0 3px {box_c}18, 0 0 18px {box_c}35;
  z-index:50;
}}
.bbox-lbl {{
  position:absolute; pointer-events:none;
  background:{box_c}; color:#fff;
  font-size:9.5px; font-weight:700; font-family:monospace;
  padding:1px 8px; border-radius:4px; white-space:nowrap;
  z-index:51; letter-spacing:0.5px;
}}
.bbox-badge {{
  position:absolute; pointer-events:none;
  background:{bdg_bg}; border:1.5px solid {box_c};
  color:{bdg_fg};
  font-size:9.5px; font-weight:800; font-family:monospace;
  padding:1px 8px; border-radius:20px; white-space:nowrap;
  z-index:51;
  box-shadow: 0 1px 4px {box_c}22;
}}
</style>
</head>
<body>
<div id="toolbar">
  <span>📄</span>
  <span class="dn">{html.escape(uploaded_name)}</span>
  <span class="fc">📍 {html.escape(fname)}</span>
  <span class="cc">conf {conf_pct}%</span>
  <span class="sp"></span>
  <span class="ll" id="ll">locating…</span>
</div>
<div id="wrap">
  <div id="page">
    {indexed_html}
  </div>
</div>
<script>
(function(){{
  const HL   = "{hl_js}";
  const FN   = "{fn_js}";
  const CONF = {conf_pct};
  const all = Array.from(document.querySelectorAll('[data-bi]'))
                   .sort((a, b) => +a.dataset.bi - +b.dataset.bi);
  let tgt = null;
  if (HL) {{
    tgt = all.find(el => {{
      const t = el.textContent.trim();
      return t.toLowerCase().includes(HL.toLowerCase()) && t.length < HL.length * 6;
    }});
    if (!tgt) {{
      tgt = all.find(el => el.textContent.toLowerCase().includes(HL.toLowerCase()));
    }}
  }}
  if (!tgt) tgt = all[0];
  if (!tgt) {{ all.forEach(e => {{ e.style.opacity = '1'; }}); return; }}
  const tidx = all.indexOf(tgt);
  const tagLabel = ['H1','H2','H3'].includes(tgt.tagName) ? 'Heading' :
                   tgt.tagName === 'TR' ? 'Table Row' : 'Paragraph';
  document.getElementById('ll').textContent =
    tagLabel + ' · block ' + tgt.dataset.bi;
  all.forEach(e => {{ e.style.opacity = '0.38'; }});
  for (let i = Math.max(0, tidx - 3); i < Math.min(all.length, tidx + 4); i++) {{
    if (all[i] !== tgt) all[i].style.opacity = '0.68';
  }}
  tgt.style.opacity = '1';
  function hlText(el, needle) {{
    if (!needle) return;
    const walk = document.createTreeWalker(el, NodeFilter.SHOW_TEXT);
    const nodes = [];
    let n;
    while ((n = walk.nextNode())) nodes.push(n);
    nodes.forEach(node => {{
      const lo = node.textContent.toLowerCase();
      const ni = lo.indexOf(needle.toLowerCase());
      if (ni === -1) return;
      try {{
        const sp = document.createElement('span');
        sp.className = 'hlm';
        const r = document.createRange();
        r.setStart(node, ni);
        r.setEnd(node, ni + needle.length);
        r.surroundContents(sp);
      }} catch (e) {{}}
    }});
  }}
  if (HL) hlText(tgt, HL);
  function drawBox() {{
    document.querySelectorAll('.bbox-box,.bbox-lbl,.bbox-badge').forEach(e => e.remove());
    const pr  = document.getElementById('page').getBoundingClientRect();
    const tr  = tgt.getBoundingClientRect();
    const pad = 7;
    const top  = tr.top  - pr.top  - pad;
    const left = tr.left - pr.left - pad;
    const w    = tr.width  + pad * 2;
    const h    = tr.height + pad * 2;
    const page = document.getElementById('page');
    const box = document.createElement('div');
    box.className = 'bbox-box';
    Object.assign(box.style, {{ top: top+'px', left: left+'px', width: w+'px', height: h+'px' }});
    page.appendChild(box);
    const lbl = document.createElement('div');
    lbl.className = 'bbox-lbl';
    lbl.textContent = '📍 ' + FN;
    Object.assign(lbl.style, {{ top: (top - 14)+'px', left: left+'px' }});
    page.appendChild(lbl);
    const bdg = document.createElement('div');
    bdg.className = 'bbox-badge';
    bdg.textContent = 'conf ' + CONF + '%';
    Object.assign(bdg.style, {{
      top:  (top - 14) + 'px',
      left: (left + w - 120) + 'px',
    }});
    page.appendChild(bdg);
  }}
  setTimeout(drawBox, 80);
  setTimeout(drawBox, 500);
  window.addEventListener('resize', drawBox);
  setTimeout(() => {{ tgt.scrollIntoView({{ behavior: 'smooth', block: 'center' }}); }}, 150);
}})();
</script>
</body>
</html>"""

    components.html(iframe, height=viewer_height, scrolling=True)


# ─────────────────────────────────────────────────────────────────────────────
# SOURCE CONTEXT RENDERER
# ─────────────────────────────────────────────────────────────────────────────

def _render_source_context(file_path: str, field: dict, current_val: str, radius: int = 2) -> str:
    try:
        from modules.word_parser import extract_word_blocks
        blocks = extract_word_blocks(file_path)
    except Exception:
        blocks = []
    if not blocks:
        return (f"<div style='color:{_LBL};font-family:monospace;font-size:11px;'>"
                f"No readable blocks found.</div>")
    orig_val  = field.get("value", current_val)
    src_block = field.get("source_block")
    hit = None
    if src_block is not None:
        for i, b in enumerate(blocks):
            if b.get("block_id") == src_block: hit = i; break
    if hit is None and orig_val:
        for i, b in enumerate(blocks):
            if orig_val.lower() in b.get("text", "").lower(): hit = i; break
    if hit is None: hit = 0
    start = max(0, hit - radius)
    end   = min(len(blocks), hit + radius + 1)
    parts = []
    for j, b in enumerate(blocks[start:end]):
        is_hit  = (start + j) == hit
        txt     = b.get("text", "")
        esc_txt = html.escape(txt)
        if is_hit and orig_val:
            pat = re.compile(re.escape(html.escape(orig_val.strip())), re.IGNORECASE)
            esc_txt = pat.sub(
                lambda m: f"<mark style='background:#fde047;color:#111827;padding:0 2px;border-radius:3px;'>{m.group(0)}</mark>",
                esc_txt,
            )
        meta_bits = []
        if b.get("para_index")  is not None: meta_bits.append(f"§ {b['para_index']+1}")
        if b.get("table_index") is not None: meta_bits.append(f"Table {b['table_index']+1}")
        if b.get("row_index")   is not None: meta_bits.append(f"Row {b['row_index']+1}")
        meta = " · ".join(meta_bits) or b.get("block_type", "block").replace("_", " ").title()
        bc   = "2px solid #ca8a04" if is_hit else f"1px solid {_BORDER}"
        bgc  = "#fffbeb"           if is_hit else _BG2
        parts.append(
            f"<div style='margin-bottom:6px;padding:8px 12px;border:{bc};border-radius:6px;background:{bgc};'>"
            f"<div style='font-size:9px;color:{_LBL};font-family:monospace;text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'>{meta}</div>"
            f"<div style='font-size:12px;font-family:monospace;color:{_TXT};line-height:1.55;white-space:pre-wrap;word-break:break-word;'>{esc_txt}</div></div>"
        )
    return "".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — FIELDS
# ─────────────────────────────────────────────────────────────────────────────

def _render_fields_tab(
    fields: list[dict], word_result: dict, uploaded_name: str, file_path: str,
) -> None:
    eds = _edits(uploaded_name)

    if not fields:
        st.markdown(_card(
            f"<div style='color:{_LBL};font-size:13px;font-family:monospace;'>⚠ No fields extracted.</div>",
            border_color="#fde68a", bg="#fffbeb",
        ), unsafe_allow_html=True)
        return

    st.markdown(_section_header("Extracted Fields",
        f"{len(fields)} field(s) extracted · click 👁 to view in document with bounding box"),
        unsafe_allow_html=True)

    _HDR = (f"font-size:10px;font-weight:700;font-family:monospace;text-transform:uppercase;"
            f"letter-spacing:1.5px;padding:6px 4px;border-bottom:1px solid {_BORDER};")
    h1, h2, h3, h4 = st.columns([2.5, 3.5, 3.5, 1.0])
    h1.markdown(f"<div style='{_HDR}color:{_LBL};'>Field Name</div>",    unsafe_allow_html=True)
    h2.markdown(f"<div style='{_HDR}color:#059669;'>Extracted</div>",    unsafe_allow_html=True)
    h3.markdown(f"<div style='{_HDR}color:#2563eb;'>Modified</div>",     unsafe_allow_html=True)
    h4.markdown(f"<div style='{_HDR}color:{_LBL};text-align:center;'>Actions</div>", unsafe_allow_html=True)
    st.markdown("<div style='height:4px;'></div>", unsafe_allow_html=True)

    _EM_KEY = _wa_key(uploaded_name, "edit_mode")
    _EY_KEY = _wa_key(uploaded_name, "eye_open")
    if _EM_KEY not in st.session_state: st.session_state[_EM_KEY] = set()
    if _EY_KEY not in st.session_state: st.session_state[_EY_KEY] = set()

    for idx, field in enumerate(fields):
        fname      = field.get("field_name", f"Field_{idx}")
        extracted  = field.get("value", "")
        conf       = float(field.get("confidence", 0.9))
        modified   = eds.get(fname, extracted)
        in_edit    = fname in st.session_state[_EM_KEY]
        eye_open   = fname in st.session_state[_EY_KEY]
        is_changed = modified != extracted

        c1, c2, c3, c4 = st.columns([2.5, 3.5, 3.5, 1.0])

        with c1:
            # NOTE: AI badge intentionally removed — no longer shown
            st.markdown(
                f"<div style='font-size:12px;font-weight:600;color:{_TXT};font-family:monospace;"
                f"padding:6px 4px 2px 4px;line-height:1.4;word-break:break-word;'>"
                f"{html.escape(fname)}</div>"
                f"<div style='padding:0 4px 6px 4px;'>{_conf_badge(conf)}</div>",
                unsafe_allow_html=True)

        with c2:
            st.markdown(
                f"<div style='font-size:12px;color:{_TXT};font-family:monospace;background:{_BG2};"
                f"border:1px solid {_BORDER};padding:7px 10px;border-radius:5px;min-height:34px;"
                f"line-height:1.5;white-space:pre-wrap;word-break:break-word;'>"
                f"{html.escape(extracted) if extracted else f'<span style=\"color:{_LBL2};\">—</span>'}</div>",
                unsafe_allow_html=True)

        with c3:
            if in_edit:
                st.text_input("modified_value", value=modified,
                              key=f"_wamv_{idx}_{fname}", label_visibility="collapsed")
            else:
                _badge = (
                    f"<span style='margin-left:6px;font-size:9px;color:#2563eb;border:1px solid #2563eb;"
                    f"border-radius:10px;padding:1px 5px;white-space:nowrap;background:#eff6ff;'>✏ edited</span>"
                    if is_changed else ""
                )
                _bgcss = (
                    f"color:{_TXT};background:#eff6ff;border:1px solid #bfdbfe;"
                    if is_changed else
                    f"color:{_TXT};background:{_BG2};border:1px solid {_BORDER};"
                )
                st.markdown(
                    f"<div style='font-size:12px;font-family:monospace;{_bgcss}"
                    f"padding:7px 10px;border-radius:5px;min-height:34px;"
                    f"line-height:1.5;white-space:pre-wrap;word-break:break-word;'>"
                    f"{html.escape(modified) if modified else f'<span style=\"color:{_LBL2};\">—</span>'}"
                    f"{_badge}</div>",
                    unsafe_allow_html=True)

        with c4:
            be, beye = st.columns(2)
            with be:
                lbl = "💾" if in_edit else "✏️"
                if st.button(lbl, key=f"_wabtn_edit_{idx}",
                             help="Save" if in_edit else "Edit", use_container_width=True):
                    if in_edit:
                        saved = st.session_state.get(f"_wamv_{idx}_{fname}", modified)
                        _sync_edit(fname, saved, uploaded_name)
                        st.session_state[_EM_KEY].discard(fname)
                    else:
                        st.session_state[_EM_KEY].add(fname)
                    st.rerun()

            with beye:
                eye_lbl  = "✕" if eye_open else "👁"
                eye_help = "Hide document view" if eye_open else "View in document with bounding box"
                if st.button(eye_lbl, key=f"_wabtn_eye_{idx}",
                             help=eye_help, use_container_width=True):
                    if eye_open:
                        st.session_state[_EY_KEY].discard(fname)
                    else:
                        st.session_state[_EY_KEY] = {fname}
                        st.session_state[_EM_KEY].discard(fname)
                    st.rerun()

        if eye_open and file_path and os.path.exists(file_path):
            st.markdown(
                f"<div style='margin:4px 0 2px 0;padding:8px 14px;background:#f0f9ff;"
                f"border:1px solid {_BORDER2};border-left:4px solid #0369a1;border-radius:8px 8px 0 0;"
                f"font-size:10px;font-weight:700;color:#0369a1;font-family:monospace;"
                f"text-transform:uppercase;letter-spacing:1.2px;display:flex;align-items:center;gap:8px;'>"
                f"<span>📄</span><span>Document Viewer — {html.escape(fname)}</span>"
                f"<span style='margin-left:auto;font-weight:400;color:{_LBL};font-size:9px;'>"
                f"bounding box · conf {int(float(field.get('confidence', 0.9)) * 100)}%</span>"
                f"</div>",
                unsafe_allow_html=True,
            )
            _render_word_document_viewer(
                file_path=file_path, field=field,
                current_val=modified, uploaded_name=uploaded_name,
                viewer_height=540,
            )

        st.markdown(
            f"<div style='height:1px;background:{_BORDER};margin:2px 0 4px 0;'></div>",
            unsafe_allow_html=True)

    # ── Add New Field ─────────────────────────────────────────────────────────
    st.markdown("<div style='margin-top:20px;'></div>", unsafe_allow_html=True)
    st.markdown(
        f"<div style='height:1px;background:linear-gradient(90deg,{_BORDER2},{_BG});margin-bottom:16px;'></div>",
        unsafe_allow_html=True)
    st.markdown(_section_header("Add New Field", "manually inject a custom key-value pair"),
                unsafe_allow_html=True)

    _ANF_KEY = _wa_key(uploaded_name, "add_field_open")
    if _ANF_KEY not in st.session_state: st.session_state[_ANF_KEY] = False

    if not st.session_state[_ANF_KEY]:
        anf_col, _ = st.columns([2, 5])
        with anf_col:
            if st.button("＋  Add New Field", key=f"_wa_anf_open_{uploaded_name}",
                         help="Manually add a custom field", use_container_width=True):
                st.session_state[_ANF_KEY] = True; st.rerun()
    else:
        st.markdown(
            f"<div style='background:{_BG2};border:1px solid {_BORDER2};"
            f"border-left:4px solid #7c3aed;border-radius:8px;padding:16px 18px;margin-bottom:10px;'>"
            f"<div style='font-size:10px;font-weight:700;color:#7c3aed;font-family:monospace;"
            f"text-transform:uppercase;letter-spacing:1.5px;margin-bottom:12px;'>✏ New Field</div>",
            unsafe_allow_html=True)
        nf1, nf2 = st.columns([1, 1])
        with nf1:
            st.markdown(f"<div style='font-size:10px;font-weight:700;color:{_LBL};font-family:monospace;"
                        f"text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'>Field Name</div>",
                        unsafe_allow_html=True)
            new_fname = st.text_input("nf_name", value="", placeholder="e.g. Policy Number",
                                      key=f"_wa_anf_name_{uploaded_name}", label_visibility="collapsed")
        with nf2:
            st.markdown(f"<div style='font-size:10px;font-weight:700;color:{_LBL};font-family:monospace;"
                        f"text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'>Field Value</div>",
                        unsafe_allow_html=True)
            new_fval = st.text_input("nf_val", value="", placeholder="e.g. POL-2024-00123",
                                     key=f"_wa_anf_val_{uploaded_name}", label_visibility="collapsed")
        st.markdown("</div>", unsafe_allow_html=True)

        bs, bc2, _ = st.columns([1.2, 1, 4.8])
        with bs:
            if st.button("💾 Save Field", key=f"_wa_anf_save_{uploaded_name}", use_container_width=True):
                fn2 = (new_fname or "").strip(); fv2 = (new_fval or "").strip()
                existing_names = {f.get("field_name", "") for f in fields}
                if not fn2:
                    st.error("Field name cannot be empty.")
                elif fn2 in existing_names:
                    st.error(f'Field "{fn2}" already exists.')
                else:
                    _wk = _wa_key(uploaded_name, "extra_fields")
                    if _wk not in st.session_state: st.session_state[_wk] = []
                    st.session_state[_wk].append({
                        "field_name": fn2, "value": fv2, "confidence": 1.0,
                        "source_block": None, "source_text": fv2,
                        "source_para": None, "source_table": None,
                        "source_row": None, "source_col": None, "_user_added": True,
                    })
                    _sync_edit(fn2, fv2, uploaded_name)
                    st.session_state[_ANF_KEY] = False
                    st.session_state.pop(f"_wa_anf_name_{uploaded_name}", None)
                    st.session_state.pop(f"_wa_anf_val_{uploaded_name}", None)
                    st.toast(f'✅ Field "{fn2}" added!'); st.rerun()
        with bc2:
            if st.button("✕ Cancel", key=f"_wa_anf_cancel_{uploaded_name}", use_container_width=True):
                st.session_state[_ANF_KEY] = False
                st.session_state.pop(f"_wa_anf_name_{uploaded_name}", None)
                st.session_state.pop(f"_wa_anf_val_{uploaded_name}", None)
                st.rerun()


def _location_label(field: dict) -> str:
    if field.get("source_table") is not None:
        lbl = f"Table {field['source_table']+1}"
        if field.get("source_row") is not None: lbl += f", Row {field['source_row']+1}"
        return lbl
    if field.get("source_para") is not None:
        return f"Paragraph {field['source_para']+1}"
    return "Document"


# ─────────────────────────────────────────────────────────────────────────────
# SUMMARY HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _build_field_kv(fields: list[dict], uploaded_name: str) -> dict[str, str]:
    eds = _edits(uploaded_name)
    return {
        f.get("field_name", ""): eds.get(f.get("field_name", ""), f.get("value", ""))
        for f in fields if f.get("field_name")
    }


def _auto_summarise_llm(fields: list[dict], doc_label: str,
                        raw_text: str, uploaded_name: str) -> str:
    kv   = _build_field_kv(fields, uploaded_name)
    hist = _edit_history(uploaded_name)

    field_lines = []
    for fname, val in kv.items():
        orig = hist.get(fname, [{}])[0].get("from", "") if fname in hist and hist[fname] else ""
        field_lines.append(
            f"  {fname}: {val}" + (f"  [was: {orig}]" if orig and orig != val else ""))
    fields_block = "\n".join(field_lines) or "(no fields extracted)"

    system = (
        f"You are a senior document analyst. Generate a concise factual summary (max 200 words) "
        f"of this {doc_label} document. Write natural prose — do NOT list field names. "
        f"Return ONLY the summary text with no preamble or markdown."
    )
    user = (
        f"Document type: {doc_label}\n\nFIELD VALUES:\n{fields_block}\n\n"
        f"DOCUMENT TEXT:\n{raw_text[:4000]}"
        + ("\n[...truncated...]" if len(raw_text) > 4000 else "")
    )

    client = _get_oai_client()
    result = _chat(client, _std_model(), system, user, max_tokens=400)
    if result:
        return re.sub(r"^```.*?```$", "", result, flags=re.DOTALL).strip()

    # Static fallback
    parts = [f"This is a **{doc_label}**."]
    for fn, lbl in [
        ("Claim Number", "claim number"), ("Policy Number", "policy number"),
        ("Insured", "insured party"), ("Claimant Name", "claimant"),
        ("Carrier", "carrier"), ("Loss Date", "loss date"),
        ("Status", "status"), ("Effective Date", "effective date"),
        ("Case Number", "case number"), ("Plaintiff", "plaintiff"),
        ("Defendant", "defendant"), ("Filing Date", "filing date"),
    ]:
        if fn in kv and kv[fn]: parts.append(f"The {lbl} is **{kv[fn]}**.")
    if raw_text and len(parts) < 3:
        parts.append(f"Excerpt: {' '.join(raw_text.split()[:60])}…")
    return " ".join(parts)


def _regenerate_summary(fields, doc_label, raw_text, uploaded_name) -> str | None:
    return _auto_summarise_llm(fields, doc_label, raw_text, uploaded_name) or None


# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — SUMMARY
# ─────────────────────────────────────────────────────────────────────────────

def _render_summary_tab(word_result: dict, fields: list[dict], uploaded_name: str) -> None:
    clf       = word_result.get("doc_classification", {})
    doc_label = clf.get("doc_type_label", "General Document")
    doc_conf  = float(clf.get("confidence", 0.0))
    raw_text  = word_result.get("raw_text", "")
    meta      = _DOC_TYPE_META.get(doc_label, _DOC_TYPE_META["General Document"])
    _SK       = _wa_key(uploaded_name, "summary_override")
    summary   = st.session_state.get(_SK, "")
    eds       = _edits(uploaded_name)
    hist      = _edit_history(uploaded_name)

    # ── Classification banner ─────────────────────────────────────────────────
    st.markdown(
        f"<div style='background:{meta['bg']};border:1px solid {meta['color']}30;"
        f"border-left:4px solid {meta['color']};border-radius:8px;"
        f"padding:14px 18px;margin-bottom:16px;'>"
        f"<div style='display:flex;align-items:center;gap:12px;'>"
        f"<span style='font-size:28px;'>{meta['icon']}</span>"
        f"<div>"
        f"<div style='font-size:20px;font-weight:800;color:{meta['color']};"
        f"font-family:monospace;text-transform:uppercase;letter-spacing:2px;'>{doc_label}</div>"
        f"<div style='font-size:11px;color:{_LBL};margin-top:2px;'>"
        f"Classification confidence: {_conf_badge(doc_conf)}</div>"
        f"</div></div></div>",
        unsafe_allow_html=True)

    changed        = [(fn, h) for fn, h in hist.items() if h]
    is_regenerated = bool(st.session_state.get(_wa_key(uploaded_name, "summary_regen_flag")))

    btn_col, status_col = st.columns([2, 6])
    with btn_col:
        regen_label    = "🔄 Re-regenerate Summary" if is_regenerated else "🔄 Regenerate with Edits"
        regen_disabled = not changed
        if st.button(
            regen_label,
            key=f"_wa_regen_summary_{uploaded_name}",
            help="Regenerate summary using edited field values" if not regen_disabled else "Make edits in the Fields tab first",
            disabled=regen_disabled,
            use_container_width=True,
        ):
            with st.spinner("Regenerating summary…"):
                ns = _regenerate_summary(fields, doc_label, raw_text, uploaded_name)
            if ns:
                st.session_state[_SK] = ns
                st.session_state[_wa_key(uploaded_name, "summary_regen_flag")] = True
                st.toast("✅ Summary regenerated!"); st.rerun()
            else:
                st.error("Could not connect to intelligence engine.")

    with status_col:
        if is_regenerated:
            st.markdown(
                f"<div style='font-size:11px;color:#16a34a;font-family:monospace;"
                f"padding-top:8px;'>✓ Showing regenerated summary · based on edited values</div>",
                unsafe_allow_html=True)
        elif changed:
            st.markdown(
                f"<div style='font-size:11px;color:#ca8a04;font-family:monospace;"
                f"padding-top:8px;'>⚠ {len(changed)} field(s) edited — click Regenerate to update summary</div>",
                unsafe_allow_html=True)

    if is_regenerated:
        if st.button("↩ Reset to auto-generated", key=f"_wa_reset_sum_{uploaded_name}"):
            st.session_state.pop(_wa_key(uploaded_name, "summary_regen_flag"), None)
            st.session_state.pop(_SK, None)
            st.rerun()

    st.markdown(_section_header("Document Summary"), unsafe_allow_html=True)

    if summary:
        annotated = summary
        if not is_regenerated:
            for fname, new_val in eds.items():
                changes = hist.get(fname, [])
                if not changes:
                    continue
                old_val = changes[0].get("from", "")
                if old_val and old_val != new_val and old_val in annotated:
                    annotated = annotated.replace(
                        old_val,
                        f"<span style='background:#eff6ff;color:#2563eb;"
                        f"border-radius:3px;padding:0 3px;font-weight:600;"
                        f"border-bottom:2px solid #2563eb;'"
                        f"title='Edited from: {old_val}'>{new_val}</span>",
                        1,
                    )

        border_color = "#16a34a" if is_regenerated else meta["color"]
        label_html   = (
            f"<div style='font-size:9px;font-weight:700;color:#16a34a;"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;"
            f"margin-bottom:6px;'>✓ Regenerated — uses your edited values</div>"
            if is_regenerated else
            f"<div style='font-size:9px;font-weight:700;color:{_LBL};"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;"
            f"margin-bottom:6px;'>⚡ Auto-generated on document load</div>"
        )
        st.markdown(
            f"<div style='background:{_BG2};border:1px solid {border_color}30;"
            f"border-radius:8px;padding:16px 20px;font-size:13px;color:{_TXT};"
            f"line-height:1.9;'>{label_html}{annotated}</div>",
            unsafe_allow_html=True)

        if changed and not is_regenerated:
            rows_html = ""
            for fname, fchanges in changed:
                old_v = fchanges[0].get("from", "—")
                new_v = eds.get(fname, fchanges[-1].get("to", "—"))
                rows_html += (
                    f"<div style='display:grid;grid-template-columns:180px 1fr auto 1fr;"
                    f"gap:8px;padding:6px 0;border-bottom:1px solid {_BORDER};align-items:center;'>"
                    f"<span style='font-size:11px;font-weight:600;color:{_TXT};"
                    f"font-family:monospace;'>{html.escape(fname)}</span>"
                    f"<span style='font-size:11px;color:{_LBL};font-family:monospace;"
                    f"text-decoration:line-through;word-break:break-word;'>{html.escape(old_v)}</span>"
                    f"<span style='font-size:13px;color:{_LBL};'>→</span>"
                    f"<span style='font-size:11px;color:#2563eb;font-family:monospace;"
                    f"font-weight:600;word-break:break-word;'>{html.escape(new_v)}</span>"
                    f"</div>"
                )
            st.markdown(
                f"<div style='background:#fffbeb;border:1px solid #fde68a;"
                f"border-radius:8px;padding:12px 16px;margin-top:12px;'>"
                f"<div style='font-size:10px;font-weight:700;color:#b45309;"
                f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;"
                f"margin-bottom:6px;'>⚠ {len(changed)} Pending Edit(s)</div>"
                f"<div style='font-size:9px;color:{_LBL};font-family:monospace;"
                f"margin-bottom:8px;'>These edits are not yet in the summary. "
                f"Click \"Regenerate with Edits\" to update it.</div>"
                f"{rows_html}</div>",
                unsafe_allow_html=True)
    else:
        st.info("Summary is being generated…")

    wc  = len(raw_text.split())
    cc  = len(raw_text)
    st.markdown(
        f"<div style='display:flex;gap:14px;margin-top:14px;flex-wrap:wrap;'>"
        + "".join(
            f"<div style='background:{_BG2};border:1px solid {_BORDER};"
            f"border-radius:6px;padding:8px 14px;'>"
            f"<div style='font-size:9px;color:{_LBL};font-family:monospace;"
            f"text-transform:uppercase;letter-spacing:1px;'>{lbl}</div>"
            f"<div style='font-size:14px;font-weight:700;color:#2563eb;"
            f"font-family:monospace;margin-top:2px;'>{val}</div></div>"
            for lbl, val in [
                ("Fields", len(fields)),
                ("Words", wc),
                ("Characters", cc),
                ("Doc Type", doc_label),
            ]
        )
        + "</div>", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — RAW JSON
# ─────────────────────────────────────────────────────────────────────────────

def _render_raw_json_tab(fields: list[dict], word_result: dict, uploaded_name: str) -> None:
    eds  = _edits(uploaded_name)
    hist = _edit_history(uploaded_name)
    clf  = word_result.get("doc_classification", {})
    kv   = _build_field_kv(fields, uploaded_name)

    st.markdown(
        _section_header(
            "Extracted Key-Value Pairs",
            f"{len(kv)} fields · modifications applied",
        ),
        unsafe_allow_html=True)

    if not kv:
        st.info("No extracted fields available.")
        return

    edited_count = sum(
        1 for fn, v in kv.items()
        if fn in eds and eds[fn] != next(
            (f.get("value", "") for f in fields if f.get("field_name") == fn), "")
    )
    if edited_count:
        st.markdown(
            f"<div style='background:#eff6ff;border:1px solid #bfdbfe;"
            f"border-radius:6px;padding:8px 14px;margin-bottom:12px;"
            f"font-size:11px;font-family:monospace;color:#2563eb;'>"
            f"✏ {edited_count} field(s) show modified values below</div>",
            unsafe_allow_html=True)

    st.code(json.dumps(kv, indent=2, ensure_ascii=False), language="json")

    changed = [(fn, h) for fn, h in hist.items() if h and fn in kv]
    if changed:
        rows_html = ""
        for fname, fchanges in changed:
            orig    = fchanges[0].get("from", "—")
            current = eds.get(fname, fchanges[-1].get("to", "—"))
            rows_html += (
                f"<div style='display:grid;grid-template-columns:180px 1fr auto 1fr;"
                f"gap:8px;padding:6px 0;border-bottom:1px solid {_BORDER};align-items:center;'>"
                f"<span style='font-size:11px;font-weight:600;color:{_TXT};"
                f"font-family:monospace;'>{html.escape(fname)}</span>"
                f"<span style='font-size:11px;color:{_LBL};font-family:monospace;"
                f"text-decoration:line-through;word-break:break-word;'>{html.escape(orig)}</span>"
                f"<span style='font-size:13px;color:{_LBL};'>→</span>"
                f"<span style='font-size:11px;color:#16a34a;font-family:monospace;"
                f"font-weight:600;word-break:break-word;'>{html.escape(current)}</span>"
                f"</div>"
            )
        with st.expander(f"📋 {len(changed)} modified field(s)"):
            st.markdown(
                f"<div style='background:{_BG};border:1px solid {_BORDER};"
                f"border-radius:8px;padding:12px 16px;'>"
                f"<div style='display:grid;grid-template-columns:180px 1fr auto 1fr;"
                f"gap:8px;padding-bottom:6px;border-bottom:1px solid {_BORDER};margin-bottom:4px;'>"
                f"<span style='font-size:9px;color:{_LBL};font-family:monospace;"
                f"text-transform:uppercase;letter-spacing:1px;'>Field</span>"
                f"<span style='font-size:9px;color:#dc2626;font-family:monospace;"
                f"text-transform:uppercase;letter-spacing:1px;'>Original</span>"
                f"<span></span>"
                f"<span style='font-size:9px;color:#16a34a;font-family:monospace;"
                f"text-transform:uppercase;letter-spacing:1px;'>Modified</span>"
                f"</div>{rows_html}</div>",
                unsafe_allow_html=True)

    full_json_str = json.dumps(kv, indent=2, ensure_ascii=False)
    st.markdown(
        f"<div style='font-size:11px;color:{_LBL};font-family:monospace;margin:10px 0;'>"
        f"⬇ {len(kv)} fields · modified values included</div>",
        unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "📥 Download JSON",
            data=full_json_str,
            file_name=f"{os.path.splitext(uploaded_name)[0]}_fields.json",
            mime="application/json",
            use_container_width=True,
        )
    with c2:
        if st.button("📋 Copy to clipboard", use_container_width=True,
                     key=f"_wa_copy_json_{uploaded_name}"):
            st.toast("Copied!")
            st.session_state["_wa_json_clipboard"] = full_json_str

    raw_text = word_result.get("raw_text", "")
    if raw_text:
        st.markdown("<div style='margin-top:16px;'></div>", unsafe_allow_html=True)
        with st.expander("📄 Full extracted text"):
            st.text_area("wa_raw_text", value=raw_text, height=300,
                         label_visibility="collapsed")
            st.download_button(
                "📥 Download raw text", data=raw_text,
                file_name=f"{os.path.splitext(uploaded_name)[0]}_text.txt",
                mime="text/plain", use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — TRANSFORMATION JOURNEY
# ─────────────────────────────────────────────────────────────────────────────

def _render_journey_tab(fields: list[dict], uploaded_name: str) -> None:
    hist          = _edit_history(uploaded_name)
    session_start = st.session_state.get("_session_start", "")
    now_str       = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    changed_fields   = [(f, hist[f.get("field_name", "")]) for f in fields
                        if f.get("field_name", "") in hist and hist[f.get("field_name", "")]]
    unchanged_fields = [f for f in fields
                        if f.get("field_name", "") not in hist
                        or not hist.get(f.get("field_name", ""), [])]
    edit_count = len(changed_fields)

    last_edit_ts = ""
    if edit_count:
        all_ts = [
            ch["timestamp"]
            for fn, chs in hist.items()
            for ch in chs
            if ch.get("timestamp")
        ]
        if all_ts:
            last_edit_ts = max(all_ts)[:19].replace("T", " ")

    st.markdown(
        f"<div style='background:{_BG2};border:1px solid {_BORDER};"
        f"border-radius:10px;padding:16px 20px;margin-bottom:20px;'>"
        f"<div style='font-size:10px;font-weight:700;color:#b45309;"
        f"font-family:monospace;text-transform:uppercase;letter-spacing:2px;"
        f"margin-bottom:14px;'>⚡ Pipeline Trace</div>"
        f"<div style='display:grid;grid-template-columns:160px 1fr;gap:8px;"
        f"padding:8px 0;border-bottom:1px solid {_BORDER};align-items:start;'>"
        f"<span style='font-size:10px;font-weight:700;color:#b45309;"
        f"font-family:monospace;text-transform:uppercase;letter-spacing:.8px;'>"
        f"📄 FILE PARSED</span>"
        f"<span style='font-size:11px;color:{_TXT2};font-family:monospace;'>"
        f"→ Fields extracted from uploaded document &nbsp;"
        f"<span style='color:{_LBL};'>"
        f"{session_start[:19].replace('T', ' ') if session_start else now_str}"
        f"</span></span></div>"
        f"<div style='display:grid;grid-template-columns:160px 1fr;gap:8px;"
        f"padding:8px 0;align-items:start;'>"
        f"<span style='font-size:10px;font-weight:700;color:#2563eb;"
        f"font-family:monospace;text-transform:uppercase;letter-spacing:.8px;'>"
        f"✏️ USER EDITS</span>"
        f"<span style='font-size:11px;color:{_TXT2};font-family:monospace;'>"
        + (
            f"→ {edit_count} field(s) manually updated &nbsp;"
            f"<span style='color:{_LBL};'>{last_edit_ts}</span>"
            if edit_count else
            f"→ <span style='color:{_LBL};'>No edits made this session</span>"
        )
        + f"</span></div></div>",
        unsafe_allow_html=True)

    st.markdown(_section_header("Field Transformation Timeline"), unsafe_allow_html=True)

    def _step_circle(n: int, color: str) -> str:
        return (
            f"<div style='width:26px;height:26px;border-radius:50%;"
            f"background:{color}15;border:2px solid {color};"
            f"display:flex;align-items:center;justify-content:center;"
            f"font-size:11px;font-weight:700;color:{color};"
            f"font-family:monospace;flex-shrink:0;'>{n}</div>"
        )

    def _field_card(field: dict, fchanges: list) -> None:
        fname     = field.get("field_name", "")
        extracted = field.get("value", "")
        is_mod    = bool(fchanges)
        border    = "#fde68a" if is_mod else _BORDER
        bg        = "#fffbeb" if is_mod else _BG
        mod_badge = (
            f"<span style='margin-left:8px;font-size:9px;font-weight:700;"
            f"color:#b45309;background:#fef9c3;border:1px solid #fde047;"
            f"border-radius:10px;padding:2px 8px;font-family:monospace;'>MODIFIED</span>"
            if is_mod else ""
        )
        step1_ts = session_start[:19].replace("T", " ") if session_start else now_str
        loc      = _location_label(field)
        src_text = field.get("source_text", "")

        card_html = (
            f"<div style='background:{bg};border:1px solid {border};"
            f"border-radius:10px;padding:16px 18px;margin-bottom:12px;'>"
            f"<div style='font-size:12px;font-weight:700;color:{_TXT};"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;"
            f"margin-bottom:14px;'>{html.escape(fname)}{mod_badge}</div>"
            f"<div style='display:flex;gap:12px;margin-bottom:10px;'>"
            f"{_step_circle(1, '#16a34a')}"
            f"<div style='flex:1;'>"
            f"<div style='display:flex;justify-content:space-between;align-items:center;"
            f"margin-bottom:4px;'>"
            f"<span style='font-size:10px;font-weight:700;color:#16a34a;"
            f"font-family:monospace;text-transform:uppercase;'>Extracted from Document</span>"
            f"<span style='font-size:9px;color:{_LBL};font-family:monospace;'>⏱ {step1_ts}</span>"
            f"</div>"
            f"<div style='font-size:11px;color:{_LBL};font-family:monospace;margin-bottom:5px;'>{html.escape(loc)}</div>"
            f"<div style='background:{_BG2};border:1px solid {_BORDER};border-radius:5px;"
            f"padding:8px 12px;font-size:12px;color:{_TXT};font-family:monospace;"
            f"word-break:break-word;min-height:32px;'>"
            f"{html.escape(extracted) if extracted else f'<span style=\"color:{_LBL2};\">—</span>'}"
            f"</div>"
            + (
                f"<div style='font-size:10px;color:{_LBL};font-family:monospace;"
                f"background:{_BG2};border-left:2px solid {_BORDER2};padding:4px 8px;"
                f"margin-top:5px;border-radius:0 4px 4px 0;font-style:italic;'>"
                f"📄 {html.escape(src_text)}</div>"
                if src_text else ""
            )
            + f"</div></div>"
            f"<div style='display:flex;gap:12px;margin-bottom:10px;'>"
            f"{_step_circle(2, '#2563eb')}"
            f"<div style='flex:1;'>"
            f"<div style='display:flex;justify-content:space-between;align-items:center;"
            f"margin-bottom:4px;'>"
            f"<span style='font-size:10px;font-weight:700;color:#2563eb;"
            f"font-family:monospace;text-transform:uppercase;'>→ Intelligence Extraction</span>"
            f"<span style='font-size:9px;color:{_LBL};font-family:monospace;'>⏱ {step1_ts} · word_intelligence</span>"
            f"</div>"
            f"<div style='font-size:11px;color:{_LBL};font-family:monospace;'>"
            f"Extracted by AI — document-type field list applied</div>"
            f"</div></div>"
        )

        for i, ch in enumerate(fchanges):
            ts     = (ch.get("timestamp", "")[:19] or "").replace("T", " ")
            from_v = ch.get("from", "")
            to_v   = ch.get("to", "")
            card_html += (
                f"<div style='display:flex;gap:12px;margin-bottom:8px;'>"
                f"{_step_circle(i + 3, '#ca8a04')}"
                f"<div style='flex:1;'>"
                f"<div style='display:flex;justify-content:space-between;align-items:center;"
                f"margin-bottom:6px;'>"
                f"<span style='font-size:10px;font-weight:700;color:#b45309;"
                f"font-family:monospace;text-transform:uppercase;'>→ User Edit</span>"
                f"<span style='font-size:9px;color:{_LBL};font-family:monospace;'>"
                f"⏱ {ts} · _sync_edit()</span>"
                f"</div>"
                f"<div style='display:flex;gap:10px;align-items:center;'>"
                f"<div style='flex:1;background:#fef2f2;border:1px solid #fecaca;"
                f"border-radius:5px;padding:7px 12px;font-size:12px;"
                f"color:#dc2626;font-family:monospace;word-break:break-word;'>"
                f"FROM: {html.escape(from_v or '—')}</div>"
                f"<span style='font-size:16px;color:{_LBL};'>→</span>"
                f"<div style='flex:1;background:#f0fdf4;border:1px solid #bbf7d0;"
                f"border-radius:5px;padding:7px 12px;font-size:12px;"
                f"color:#16a34a;font-family:monospace;word-break:break-word;'>"
                f"TO: {html.escape(to_v or '—')}</div>"
                f"</div></div></div>"
            )

        card_html += "</div>"
        st.markdown(card_html, unsafe_allow_html=True)

    for field, chs in changed_fields:
        _field_card(field, chs)

    if unchanged_fields:
        with st.expander(f"📋 {len(unchanged_fields)} unchanged field(s)"):
            for f in unchanged_fields:
                _field_card(f, hist.get(f.get("field_name", ""), []))


# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — SIGNALS
# ─────────────────────────────────────────────────────────────────────────────

def _classify_signal_severity(sig: dict) -> str:
    raw_sev = str(sig.get("severity", "")).strip().title()
    if raw_sev in _TAXONOMY:
        return raw_sev
    _LEGACY = {
        "severely_high": "Highly Severe",
        "moderate":      "Moderate",
        "low":           "Low",
    }
    legacy = _LEGACY.get(raw_sev.lower().replace(" ", "_"))
    if legacy:
        return legacy
    desc = (sig.get("description", "") + " " + sig.get("evidence", "")).lower()
    hs   = ["fatal", "death", "fatality", "catastrophic", "permanent disab",
            "punitive", "class action", "multi-party", "fraud confirmed"]
    if any(k in desc for k in hs):
        return "Highly Severe"
    if any(k in desc for k in ["legal", "litigation", "lawsuit", "severe"]):
        return "High"
    if any(k in desc for k in ["missing", "gap", "inconsistent", "risk"]):
        return "Moderate"
    return "Low"


def _rule_based_signals(kv: dict[str, str], doc_label: str, raw_text: str) -> list[dict]:
    """
    Deterministic rule-based signal detection — runs without any LLM.
    Covers date anomalies, missing critical fields, suspicious values,
    legal/financial red flags, and document-type-specific checks.
    Always returns a list (may be empty if document is clean).
    """
    import datetime as _dt

    signals: list[dict] = []
    lower = raw_text.lower()

    def _sig(signal: str, severity: str, description: str, evidence: str = "") -> None:
        signals.append({"signal": signal, "severity": severity,
                        "description": description, "evidence": evidence})

    # ── 1. Date anomaly checks ────────────────────────────────────────────────
    today = _dt.date.today()
    _DATE_FIELDS = [
        "Filing Date", "FILING DATE", "Effective Date", "EFFECTIVE DATE",
        "Expiration Date", "EXPIRATION DATE", "Loss Date", "Date of Loss",
        "Invoice Date", "Due Date", "Date of Service",
    ]
    _MONTH_MAP = {
        "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
        "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    }

    def _parse_date(s: str):
        s = s.strip()
        # Try ISO / common formats
        for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y",
                    "%b %d, %Y", "%b. %d, %Y", "%B %d %Y"):
            try:
                return _dt.datetime.strptime(s, fmt).date()
            except ValueError:
                pass
        # "Feb. 29th, 2024" / "Sept. 25th, 2025" — strip ordinal suffix
        m = re.match(
            r"([A-Za-z]+)\.?\s+(\d{1,2})(?:st|nd|rd|th)?,?\s+(\d{4})", s, re.IGNORECASE)
        if m:
            mon_str = m.group(1).lower()[:3]
            mon = _MONTH_MAP.get(mon_str)
            if mon:
                try:
                    return _dt.date(int(m.group(3)), mon, int(m.group(2)))
                except ValueError:
                    pass
        return None

    for fld in _DATE_FIELDS:
        val = kv.get(fld, "")
        if not val:
            continue
        d = _parse_date(val)
        if d is None:
            continue
        # Future date check
        if d > today:
            days_ahead = (d - today).days
            if fld.lower() in ("filing date", "filing date"):
                _sig("Filing Date in Future",
                     "Highly Severe",
                     f"The filing date ({val}) is set in the future ({days_ahead} days ahead), "
                     f"which is not permissible in legal proceedings.",
                     f"{fld}: {val}")
            elif "effective" in fld.lower():
                _sig("Effective Date in Future",
                     "High",
                     f"The effective date ({val}) is set in the future, raising concerns "
                     f"about the validity of the document timeline.",
                     f"{fld}: {val}")
            elif "expir" in fld.lower():
                pass  # future expiry is expected
            else:
                _sig(f"{fld} is a Future Date",
                     "Moderate",
                     f"The field '{fld}' contains a future date ({val}), which may indicate "
                     f"a data entry error.",
                     f"{fld}: {val}")
        # Invalid calendar date (e.g. Feb 29 in a non-leap year) — already caught by _parse_date
        # Very old date (>30 years)
        if d.year < today.year - 30:
            _sig(f"Unusually Old Date — {fld}",
                 "Moderate",
                 f"'{fld}' contains a date ({val}) that is over 30 years in the past, "
                 f"which may indicate a data error.",
                 f"{fld}: {val}")

    # Detect Feb 29 in non-leap year via raw text scan
    feb29 = re.findall(r"[Ff]eb(?:ruary|\.)?\.?\s+29(?:th)?", raw_text)
    for match in feb29:
        # Find year context
        ctx_match = re.search(
            r"[Ff]eb(?:ruary|\.)?\.?\s+29(?:th)?,?\s+(\d{4})", raw_text)
        if ctx_match:
            yr = int(ctx_match.group(1))
            import calendar as _cal
            if not _cal.isleap(yr):
                _sig("Invalid Date — February 29 in Non-Leap Year",
                     "Highly Severe",
                     f"February 29, {yr} does not exist ({yr} is not a leap year). "
                     f"This is an invalid date and may invalidate the document.",
                     ctx_match.group(0))

    # ── 2. Missing critical fields by document type ───────────────────────────
    _REQUIRED: dict[str, list[tuple[str, str]]] = {
        "Legal Document": [
            ("Case Number",  "Highly Severe", "Case Number is missing — this field is required to identify the legal proceeding."),
            ("Plaintiff",    "High",          "Plaintiff name is absent, making it impossible to identify the initiating party."),
            ("Defendant",    "High",          "Defendant name is absent; the opposing party cannot be identified."),
            ("Filing Date",  "High",          "No filing date found; the timeline of the case cannot be established."),
            ("Court",        "Moderate",      "Court / jurisdiction information is missing from the document."),
        ],
        "Insurance Policy": [
            ("Policy Number", "Highly Severe", "Policy Number is missing — the policy cannot be uniquely identified."),
            ("Insured",       "High",          "Insured name is absent from the policy."),
            ("Effective Date","High",          "No effective date found; coverage start cannot be determined."),
            ("Expiration Date","High",         "No expiration date found; coverage end cannot be determined."),
            ("Premium Amount","Moderate",      "Premium amount is not captured in the extracted fields."),
        ],
        "Claims Report": [
            ("Claim Number",  "Highly Severe", "Claim Number is missing — the claim cannot be identified."),
            ("Claimant Name", "High",          "Claimant name is absent from the report."),
            ("Loss Date",     "High",          "Date of loss is not present; claim timeline cannot be established."),
            ("Claim Status",  "Moderate",      "Claim status is not captured; current state is unknown."),
        ],
        "Medical Report": [
            ("Patient Name",  "Highly Severe", "Patient name is missing from the medical report."),
            ("Diagnosis",     "High",          "No diagnosis field found; clinical purpose cannot be determined."),
            ("Date of Service","High",         "Date of service is absent; medical timeline cannot be established."),
        ],
        "Invoice / Bill": [
            ("Invoice Number","Highly Severe", "Invoice number is missing; the invoice cannot be uniquely referenced."),
            ("Total Amount",  "High",          "Total amount is absent; financial obligation cannot be determined."),
            ("Due Date",      "Moderate",      "Payment due date is not present."),
        ],
    }
    required = _REQUIRED.get(doc_label, [])
    # Normalise kv keys for matching
    kv_lower = {k.lower(): v for k, v in kv.items()}
    for field_name, severity, description in required:
        if not kv_lower.get(field_name.lower()):
            _sig(f"Missing Field — {field_name}", severity, description, "")

    # ── 3. Policy number in legal document (cross-reference check) ────────────
    if doc_label == "Legal Document":
        pol = kv_lower.get("policy number", "")
        if pol and pol.lower() not in lower:
            _sig("Policy Number Not Found in Body Text",
                 "Moderate",
                 f"The extracted policy number ({pol}) does not appear verbatim in the "
                 f"document body, suggesting a possible extraction error.",
                 f"Policy Number: {pol}")

    # ── 4. Conflict: filing date vs effective date ────────────────────────────
    filing_str    = kv_lower.get("filing date", "")
    effective_str = kv_lower.get("effective date", "")
    if filing_str and effective_str:
        fd = _parse_date(filing_str)
        ed = _parse_date(effective_str)
        if fd and ed and ed < fd:
            _sig("Effective Date Precedes Filing Date",
                 "High",
                 f"The effective date ({effective_str}) is earlier than the filing date "
                 f"({filing_str}), which may indicate a retroactive or erroneous entry.",
                 f"Filing Date: {filing_str} | Effective Date: {effective_str}")

    # ── 5. Legal-specific text signals ───────────────────────────────────────
    if doc_label in ("Legal Document", "General Document"):
        if re.search(r"fraud|fraudulent|misrepresent", lower):
            _sig("Fraud or Misrepresentation Allegation",
                 "Highly Severe",
                 "The document contains language referencing fraud or misrepresentation, "
                 "indicating serious legal risk.",
                 re.search(r".{0,60}(?:fraud|fraudulent|misrepresent).{0,60}", lower).group(0).strip())
        if re.search(r"punitive\s+damage", lower):
            _sig("Punitive Damages Claim",
                 "Highly Severe",
                 "A claim for punitive damages is present, indicating potential for "
                 "significantly elevated financial exposure.",
                 re.search(r".{0,60}punitive\s+damage.{0,60}", lower).group(0).strip())
        if re.search(r"class\s+action", lower):
            _sig("Class Action Exposure",
                 "Highly Severe",
                 "The document references a class action, suggesting widespread liability "
                 "across multiple claimants.",
                 re.search(r".{0,60}class\s+action.{0,60}", lower).group(0).strip())
        if re.search(r"unauthorized\s+use|no\s+valid\s+(?:license|licence|operator)", lower):
            _sig("Unauthorized Use / Unlicensed Operator",
                 "High",
                 "The document mentions unauthorized vehicle use or operation without a valid "
                 "license, which may affect coverage under the policy.",
                 re.search(r".{0,80}(?:unauthorized\s+use|no\s+valid\s+(?:license|licence|operator)).{0,80}", lower).group(0).strip())
        if re.search(r"negligence|negligent", lower):
            _sig("Negligence Claim Present",
                 "High",
                 "Negligence allegations are present in the document, indicating potential "
                 "tort liability.",
                 re.search(r".{0,60}negligen.{0,60}", lower).group(0).strip())
        if re.search(r"declaratory\s+judgment|declaratory\s+relief", lower):
            _sig("Declaratory Judgment Action",
                 "Moderate",
                 "A declaratory judgment action is present, meaning a court is being asked to "
                 "determine rights or obligations under a policy or contract.",
                 re.search(r".{0,60}declaratory\s+(?:judgment|relief).{0,60}", lower).group(0).strip())

    # ── 6. Insurance-specific signals ────────────────────────────────────────
    if doc_label in ("Insurance Policy", "Claims Report", "Loss Run Report",
                     "Certificate of Insurance", "General Document"):
        if re.search(r"coverage\s+(?:exclusion|excluded|denied|lapse)", lower):
            _sig("Coverage Exclusion or Denial",
                 "High",
                 "The document references a coverage exclusion or denial, which may leave "
                 "a party unprotected.",
                 re.search(r".{0,80}coverage\s+(?:exclusion|excluded|denied|lapse).{0,80}", lower).group(0).strip())

    # ── 7. Generic quality signals ────────────────────────────────────────────
    if len(kv) < 3:
        _sig("Very Few Fields Extracted",
             "Moderate",
             f"Only {len(kv)} field(s) were extracted from this document. The document may "
             f"be poorly structured, scanned, or contain primarily unstructured text.",
             "")

    total_words = len(raw_text.split())
    if total_words < 30:
        _sig("Document May Be Empty or Truncated",
             "High",
             f"The document contains very little text ({total_words} words), suggesting it "
             f"may be empty, image-only, or improperly parsed.",
             "")

    # ── Sort and return ───────────────────────────────────────────────────────
    order = {"Highly Severe": 0, "High": 1, "Moderate": 2, "Low": 3}
    signals.sort(key=lambda x: order.get(x["severity"], 99))
    return signals


def _run_signals_analysis(fields: list[dict], doc_label: str,
                          raw_text: str, uploaded_name: str) -> list[dict]:
    """
    Run signal detection. Always returns a list (never None).
    Tries LLM first; falls back to rule-based detection if LLM is
    unavailable or returns invalid data.
    """
    kv           = _build_field_kv(fields, uploaded_name)
    fields_block = "\n".join(f"  {k}: {v}" for k, v in kv.items()) or "(no fields)"

    system = (
        "You are a document risk and quality analyst. "
        "Given a document's extracted fields and text, identify all notable signals — "
        "risks, anomalies, gaps, red flags, or quality concerns. "
        "Categorise each signal with exactly one of these four severity levels: "
        "Highly Severe, High, Moderate, or Low.\n\n"
        "Return ONLY a JSON array, no markdown, no preamble:\n"
        '[{"signal":"<concise title>","severity":"Highly Severe|High|Moderate|Low",'
        '"description":"<1-2 sentences>","evidence":"<quoted snippet from doc or field value>"}]'
    )
    user = (
        f"Document type: {doc_label}\n\n"
        f"Extracted fields:\n{fields_block}\n\n"
        f"Document text:\n{raw_text[:4000]}"
        + ("\n[...truncated...]" if len(raw_text) > 4000 else "")
    )

    client = _get_oai_client()
    raw    = _chat(client, _std_model(), system, user, max_tokens=1200)
    parsed = _parse_json_response(raw)

    if isinstance(parsed, list) and parsed:
        valid = []
        for item in parsed:
            sev = _classify_signal_severity(item)
            valid.append({
                "signal":      str(item.get("signal", "Signal")).strip(),
                "severity":    sev,
                "description": str(item.get("description", "")).strip(),
                "evidence":    str(item.get("evidence", "")).strip(),
            })
        order = {"Highly Severe": 0, "High": 1, "Moderate": 2, "Low": 3}
        valid.sort(key=lambda x: order.get(x["severity"], 99))
        return valid

    # ── LLM unavailable or returned empty/invalid — use rule-based fallback ──
    return _rule_based_signals(kv, doc_label, raw_text)


def _render_signals_tab(fields: list[dict], word_result: dict, uploaded_name: str) -> None:
    clf       = word_result.get("doc_classification", {})
    doc_label = clf.get("doc_type_label", "General Document")
    raw_text  = word_result.get("raw_text", "")
    _SK       = _wa_key(uploaded_name, "signals_result")
    signals   = st.session_state.get(_SK)

    # ── AUTO-RUN on first visit (no button click needed) ──────────────────────
    # _run_signals_analysis always returns a list (LLM or rule-based fallback)
    _ran_key = _wa_key(uploaded_name, "signals_autorun_done")
    if signals is None and not st.session_state.get(_ran_key):
        st.session_state[_ran_key] = True
        with st.spinner("Analysing signals…"):
            result = _run_signals_analysis(fields, doc_label, raw_text, uploaded_name)
        st.session_state[_SK] = result
        signals = result

    sig_count = len(signals) if signals else 0
    st.markdown(
        _section_header("Signal Detection", f"{sig_count} signal(s) detected"),
        unsafe_allow_html=True)

    # ── Intro banner ──────────────────────────────────────────────────────────
    st.markdown(
        f"<div style='background:#f0f9ff;border:1px solid #bae6fd;"
        f"border-left:4px solid #0284c7;border-radius:8px;"
        f"padding:14px 18px;margin-bottom:18px;'>"
        f"<div style='font-size:13px;font-weight:600;color:#0369a1;margin-bottom:4px;'>"
        f"🔍 Intelligent Signal Detection</div>"
        f"<div style='font-size:12px;color:#0c4a6e;line-height:1.7;'>"
        f"Scans the document for risk indicators, data anomalies, missing information, "
        f"and quality concerns. Each signal is rated "
        f"<b style='color:#dc2626;'>Highly Severe</b>, "
        f"<b style='color:#ea580c;'>High</b>, "
        f"<b style='color:#ca8a04;'>Moderate</b>, or "
        f"<b style='color:#16a34a;'>Low</b>.</div></div>",
        unsafe_allow_html=True)

    # ── Re-run button (only shown after first run) ────────────────────────────
    if signals is not None:
        rerun_col, _ = st.columns([2, 5])
        with rerun_col:
            if st.button("🔄 Re-run Analysis", key=f"_wa_run_sig_{uploaded_name}",
                         use_container_width=True):
                with st.spinner("Re-analysing signals…"):
                    result = _run_signals_analysis(fields, doc_label, raw_text, uploaded_name)
                st.session_state[_SK] = result
                signals = result
                st.toast(f"✅ {len(result)} signal(s) detected!")
                st.rerun()

    if signals is None:
        # Should not happen now that _run_signals_analysis always returns a list,
        # but guard defensively in case of unexpected early exit
        st.info("Signal analysis loading…")
        return

    # ── Group by taxonomy ─────────────────────────────────────────────────────
    grouped: dict[str, list[dict]] = {lv: [] for lv in _TAXONOMY}
    for sig in signals:
        grouped[sig["severity"]].append(sig)

    if not signals:
        st.markdown(
            _card(
                f"<div style='color:#16a34a;font-size:13px;font-family:monospace;'>"
                f"✓ No significant signals detected — document appears clean.</div>",
                border_color="#bbf7d0", bg="#f0fdf4",
            ),
            unsafe_allow_html=True)
        return

    # ── Summary pills ─────────────────────────────────────────────────────────
    pills = "".join(
        f"<span style='background:{_TAXONOMY[lv]['bg']};"
        f"border:1px solid {_TAXONOMY[lv]['color']}40;border-radius:20px;"
        f"padding:4px 12px;font-size:11px;font-weight:700;"
        f"color:{_TAXONOMY[lv]['color']};font-family:monospace;'>"
        f"{_TAXONOMY[lv]['icon']} {lv} ({len(sigs)})</span>"
        for lv, sigs in grouped.items()
        if sigs
    )
    if pills:
        st.markdown(
            f"<div style='display:flex;flex-wrap:wrap;gap:8px;margin-bottom:18px;'>"
            f"{pills}</div>",
            unsafe_allow_html=True)

    # ── Per-group signal cards ─────────────────────────────────────────────────
    for level in ["Highly Severe", "High", "Moderate", "Low"]:
        group_sigs = grouped.get(level, [])
        if not group_sigs:
            continue
        tax = _TAXONOMY[level]
        tc  = tax["color"]

        st.markdown(
            f"<div style='display:flex;align-items:center;gap:8px;margin:16px 0 8px 0;'>"
            f"<span style='font-size:16px;'>{tax['icon']}</span>"
            f"<span style='font-size:12px;font-weight:700;color:{tc};"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1.2px;'>{level}</span>"
            f"<div style='flex:1;height:1px;background:{tc}30;'></div>"
            f"<span style='font-size:10px;color:{tc};font-family:monospace;"
            f"background:{tax['bg']};border:1px solid {tc}30;border-radius:10px;"
            f"padding:1px 8px;'>{len(group_sigs)} signal(s)</span>"
            f"</div>",
            unsafe_allow_html=True)

        for sig in group_sigs:
            ev = sig.get("evidence", "").strip()
            st.markdown(
                f"<div style='background:{_BG};border:1px solid {_BORDER};"
                f"border-left:4px solid {tc};border-radius:8px;"
                f"padding:12px 16px;margin-bottom:8px;'>"
                f"<div style='display:flex;align-items:center;gap:8px;margin-bottom:6px;'>"
                f"<span style='font-size:14px;'>{tax['icon']}</span>"
                f"<span style='font-size:12px;font-weight:700;color:{tc};"
                f"font-family:monospace;'>{html.escape(sig['signal'])}</span>"
                f"<span style='margin-left:auto;font-size:9px;color:{tc};"
                f"background:{tax['bg']};border:1px solid {tc}30;border-radius:10px;"
                f"padding:1px 7px;font-family:monospace;white-space:nowrap;'>"
                f"{tax['icon']} {level}</span>"
                f"</div>"
                f"<div style='font-size:13px;color:{_TXT};line-height:1.7;margin-bottom:6px;'>"
                f"{html.escape(sig['description'])}</div>"
                + (
                    f"<div style='font-size:11px;color:{_LBL};font-family:monospace;"
                    f"background:{_BG2};border-left:2px solid {_BORDER2};padding:5px 10px;"
                    f"border-radius:0 4px 4px 0;font-style:italic;'>"
                    f"📄 \"{html.escape(ev)}\"</div>"
                    if ev else ""
                )
                + "</div>",
                unsafe_allow_html=True)

    st.markdown("<div style='margin-top:16px;'></div>", unsafe_allow_html=True)
    _, cc = st.columns([5, 2])
    with cc:
        if st.button("🗑 Clear results", key=f"_wa_clr_sig_{uploaded_name}",
                     use_container_width=True):
            st.session_state.pop(_SK, None)
            st.session_state.pop(_wa_key(uploaded_name, "signals_autorun_done"), None)
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# TAB 6 — AI ASSISTANT
# ─────────────────────────────────────────────────────────────────────────────

def _render_ai_assistant_dimension(
    icon: str,
    title: str,
    data: dict,
    color: str,
) -> None:
    """
    Render a single AI Assistant dimension card.
    Safely handles incorrect_fields items that may be str or dict.
    """
    score    = data.get("score", 0)
    verdict  = data.get("verdict", "Review")
    findings = data.get("findings", "")

    st.markdown(
        f"<div style='background:{_BG};border:1px solid {_BORDER};"
        f"border-left:4px solid {color};border-radius:10px;"
        f"padding:18px 20px;margin-bottom:14px;'>"
        f"<div style='display:flex;align-items:center;gap:10px;margin-bottom:4px;'>"
        f"<span style='font-size:20px;'>{icon}</span>"
        f"<span style='font-size:13px;font-weight:700;color:{_TXT};"
        f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;'>{title}</span>"
        f"<div style='margin-left:auto;display:flex;align-items:center;gap:8px;'>"
        f"{_score_badge(score)}{_verdict_badge(verdict)}</div></div>"
        f"<div style='height:1px;background:{_BORDER};margin:10px 0;'></div>"
        f"<div style='font-size:13px;color:{_TXT2};line-height:1.8;'>{findings}</div>",
        unsafe_allow_html=True)

    issues    = data.get("issues", [])
    missed    = data.get("missed_fields", []) or data.get("gaps", [])
    incorrect = data.get("incorrect_fields", [])
    false_pos = data.get("false_positives", [])

    if missed:
        items_html = "".join(
            f"<span style='background:#fef2f2;border:1px solid #fecaca;"
            f"border-radius:4px;padding:2px 8px;font-size:11px;"
            f"color:#dc2626;font-family:monospace;margin:2px;'>{html.escape(str(m))}</span>"
            for m in missed
        )
        st.markdown(
            f"<div style='margin-top:8px;'>"
            f"<div style='font-size:9px;font-weight:700;color:{_LBL};"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'>"
            f"Missing / Not Detected</div>"
            f"<div style='display:flex;flex-wrap:wrap;gap:4px;'>{items_html}</div></div>",
            unsafe_allow_html=True)

    if incorrect:
        rows = []
        for item in incorrect:
            # ── FIX: handle both dict and plain string items ──────────────────
            if isinstance(item, dict):
                f_name   = str(item.get("field", ""))
                f_extr   = str(item.get("extracted", ""))
                f_expect = str(item.get("expected", ""))
            else:
                # item is a plain string description
                f_name   = str(item)
                f_extr   = ""
                f_expect = ""
            rows.append(
                f"<div style='display:grid;grid-template-columns:140px 1fr auto 1fr;"
                f"gap:8px;padding:5px 0;border-bottom:1px solid {_BORDER};align-items:start;'>"
                f"<span style='font-size:11px;font-weight:600;color:{_TXT};"
                f"font-family:monospace;word-break:break-word;'>{html.escape(f_name)}</span>"
                f"<span style='font-size:11px;color:#dc2626;font-family:monospace;"
                f"word-break:break-word;text-decoration:line-through;'>{html.escape(f_extr)}</span>"
                f"<span style='color:{_LBL};font-size:12px;'>→</span>"
                f"<span style='font-size:11px;color:#16a34a;font-family:monospace;"
                f"word-break:break-word;'>{html.escape(f_expect)}</span>"
                f"</div>"
            )
        st.markdown(
            f"<div style='margin-top:10px;background:{_BG2};border:1px solid {_BORDER};"
            f"border-radius:6px;padding:10px 14px;'>"
            f"<div style='font-size:9px;font-weight:700;color:{_LBL};"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;"
            f"margin-bottom:8px;'>Incorrect Extractions</div>{''.join(rows)}</div>",
            unsafe_allow_html=True)

    if false_pos:
        items_html = "".join(
            f"<span style='background:#fffbeb;border:1px solid #fde68a;"
            f"border-radius:4px;padding:2px 8px;font-size:11px;"
            f"color:#b45309;font-family:monospace;margin:2px;'>{html.escape(str(fp))}</span>"
            for fp in false_pos
        )
        st.markdown(
            f"<div style='margin-top:8px;'>"
            f"<div style='font-size:9px;font-weight:700;color:{_LBL};"
            f"font-family:monospace;text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;'>"
            f"Potential False Positives</div>"
            f"<div style='display:flex;flex-wrap:wrap;gap:4px;'>{items_html}</div></div>",
            unsafe_allow_html=True)

    if issues:
        rows_html = "".join(
            f"<li style='font-size:12px;color:{_TXT2};margin-bottom:4px;"
            f"font-family:monospace;'>{html.escape(str(iss))}</li>"
            for iss in issues
        )
        st.markdown(
            f"<ul style='margin:10px 0 0 16px;padding:0;list-style:disc;'>{rows_html}</ul>",
            unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)


def _run_ai_assistant(
    fields:        list[dict],
    doc_label:     str,
    raw_text:      str,
    uploaded_name: str,
    summary:       str,
    signals:       list[dict] | None,
) -> dict | None:
    kv   = _build_field_kv(fields, uploaded_name)
    hist = _edit_history(uploaded_name)

    edits_block = ""
    for fn, changes in hist.items():
        if changes:
            for ch in changes:
                ts  = ch.get("timestamp", "")[:19].replace("T", " ")
                fro = ch.get("from", "—")
                to  = ch.get("to", "—")
                edits_block += f"  [{ts}] {fn}: '{fro}' → '{to}'\n"
    if not edits_block:
        edits_block = "  (no user edits recorded this session)"

    signals_block = ""
    if signals:
        for s in signals:
            signals_block += f"  [{s['severity'].upper()}] {s['signal']}: {s['description']}\n"
    else:
        signals_block = "  (signal analysis not yet run)"

    system = (
        "You are a senior document intelligence auditor. "
        "Given a complete analysis payload for a document (fields, summary, signals, "
        "raw JSON, and edit history), produce a comprehensive quality assessment. "
        "Be specific — cite field names and evidence where relevant. "
        "For incorrect_fields, each item MUST be a JSON object with keys "
        "'field', 'extracted', and 'expected' — never a plain string. "
        "Return ONLY the following JSON structure — no markdown, no preamble:\n"
        "{\n"
        '  "overall_validation": {"score":<0-100>,"verdict":"<Pass|Review|Fail>",'
        '"summary":"<2-3 sentences>","recommended_actions":["<action>"]},\n'
        '  "extraction_accuracy": {"score":<0-100>,"verdict":"<string>","findings":"<string>",'
        '"missed_fields":[],"incorrect_fields":[{"field":"","extracted":"","expected":""}]},\n'
        '  "signal_credibility": {"score":<0-100>,"verdict":"<string>","findings":"<string>",'
        '"false_positives":[]},\n'
        '  "coverage_analysis": {"score":<0-100>,"verdict":"<string>","findings":"<string>","gaps":[]}\n'
        "}"
    )
    user = (
        f"Document type: {doc_label}\n\n"
        f"=== EXTRACTED FIELDS (JSON) ===\n{json.dumps(kv, indent=2, ensure_ascii=False)}\n\n"
        f"=== GENERATED SUMMARY ===\n{summary or '(not generated)'}\n\n"
        f"=== SIGNALS ===\n{signals_block}\n\n"
        f"=== TRANSFORMATION / EDIT HISTORY ===\n{edits_block}\n\n"
        f"=== RAW DOCUMENT TEXT (excerpt) ===\n{raw_text[:3000]}"
        + ("\n[...truncated...]" if len(raw_text) > 3000 else "")
    )

    client = _get_oai_client()
    raw    = _chat(client, _adv_model(), system, user, max_tokens=1800)
    return _parse_json_response(raw)


def _render_ai_assistant_tab(fields: list[dict], word_result: dict, uploaded_name: str) -> None:
    clf       = word_result.get("doc_classification", {})
    doc_label = clf.get("doc_type_label", "General Document")
    raw_text  = word_result.get("raw_text", "")
    _VK       = _wa_key(uploaded_name, "ai_assistant_result")
    _SK_SUM   = _wa_key(uploaded_name, "summary_override")
    _SK_SIG   = _wa_key(uploaded_name, "signals_result")
    existing  = st.session_state.get(_VK)
    summary   = st.session_state.get(_SK_SUM, "")
    signals   = st.session_state.get(_SK_SIG)

    st.markdown(_section_header("AI Assistant — Document Validation"), unsafe_allow_html=True)

    st.markdown(
        f"<div style='background:#f0f9ff;border:1px solid #bae6fd;"
        f"border-left:4px solid #0284c7;border-radius:8px;"
        f"padding:14px 18px;margin-bottom:18px;'>"
        f"<div style='font-size:13px;font-weight:600;color:#0369a1;margin-bottom:4px;'>"
        f"✅ Deep Validation</div>"
        f"<div style='font-size:12px;color:#0c4a6e;line-height:1.7;'>"
        f"This validation evaluates extraction accuracy, signal credibility, and coverage "
        f"completeness using an enhanced reasoning pass over the extracted data. "
        f"Run on demand — results are cached for this session.</div>"
        f"<div style='display:flex;gap:6px;flex-wrap:wrap;margin-top:10px;'>"
        + "".join(
            f"<span style='background:#ede9fe;border:1px solid #c4b5fd;border-radius:4px;"
            f"padding:2px 8px;font-size:10px;color:#7c3aed;font-family:monospace;'>{t}</span>"
            for t in ["📍 Fields", "📝 Summary", "🔍 Signals", "📄 JSON", "🔄 Journey"]
        )
        + "</div></div>",
        unsafe_allow_html=True)

    missing = []
    if not summary: missing.append("Summary (auto-generated on load)")
    if signals is None: missing.append("Signals (run Signal Analysis first)")
    if missing:
        st.markdown(
            f"<div style='background:#fffbeb;border:1px solid #fde68a;border-radius:6px;"
            f"padding:10px 14px;margin-bottom:14px;font-size:12px;"
            f"color:#92400e;font-family:monospace;'>"
            f"⚠ For best results, also run: {' · '.join(missing)}</div>",
            unsafe_allow_html=True)

    btn_label = "🔄 Re-run Validation" if existing else "▶ Run Validation"
    run_col, _ = st.columns([2, 5])
    with run_col:
        run_clicked = st.button(btn_label, key=f"_wa_run_ai_{uploaded_name}",
                                use_container_width=True)

    if run_clicked:
        if not fields:
            st.error("No fields to assess."); return
        with st.spinner("Running intelligence audit across all tabs…"):
            result = _run_ai_assistant(
                fields, doc_label, raw_text, uploaded_name, summary, signals)
        if result:
            st.session_state[_VK] = result
            existing = result
            st.toast("✅ Validation complete!"); st.rerun()
        else:
            st.error("Validation failed — intelligence engine unavailable.")
            return

    if not existing:
        st.markdown(
            f"<div style='background:{_BG2};border:1px solid {_BORDER};"
            f"border-radius:8px;padding:24px;text-align:center;"
            f"color:{_LBL};font-family:monospace;font-size:12px;'>"
            f"Click <strong>▶ Run Validation</strong> to start the cross-tab audit.</div>",
            unsafe_allow_html=True)
        return

    # ── Overall score banner ──────────────────────────────────────────────────
    overall    = existing.get("overall_validation", {})
    ov_score   = overall.get("score", 0)
    ov_verdict = overall.get("verdict", "Review")
    ov_summary = overall.get("summary", "")
    ov_color   = "#16a34a" if ov_score >= 80 else "#ca8a04" if ov_score >= 60 else "#dc2626"
    ov_bg      = "#f0fdf4" if ov_score >= 80 else "#fffbeb" if ov_score >= 60 else "#fef2f2"

    st.markdown(
        f"<div style='background:{ov_bg};border:2px solid {ov_color}30;"
        f"border-radius:12px;padding:20px 24px;margin-bottom:22px;'>"
        f"<div style='display:flex;align-items:center;gap:16px;margin-bottom:10px;'>"
        f"<div style='text-align:center;min-width:70px;'>"
        f"<div style='font-size:48px;font-weight:900;color:{ov_color};"
        f"font-family:monospace;line-height:1;'>{ov_score}</div>"
        f"<div style='font-size:10px;color:{_LBL};font-family:monospace;'>/ 100</div>"
        f"</div>"
        f"<div style='flex:1;'>"
        f"<div style='display:flex;align-items:center;gap:8px;margin-bottom:6px;'>"
        f"<span style='font-size:16px;font-weight:800;color:{_TXT};'>Overall Validation Score</span>"
        f"{_verdict_badge(ov_verdict)}</div>"
        f"<div style='font-size:13px;color:{_TXT2};line-height:1.7;'>{ov_summary}</div>"
        f"</div></div>"
        f"<div style='height:6px;background:{_BORDER};border-radius:3px;overflow:hidden;'>"
        f"<div style='height:100%;width:{ov_score}%;background:{ov_color};"
        f"border-radius:3px;transition:width 0.5s;'></div></div></div>",
        unsafe_allow_html=True)

    # ── Per-dimension score pills ─────────────────────────────────────────────
    dims = [
        ("extraction_accuracy", "🎯", "Extraction",  "#2563eb"),
        ("signal_credibility",  "⚡", "Signals",     "#7c3aed"),
        ("coverage_analysis",   "📋", "Coverage",    "#059669"),
    ]
    pills_html = ""
    for key, icon, label, color in dims:
        d  = existing.get(key, {})
        # Guard: d might be a string if LLM returned malformed data
        if not isinstance(d, dict):
            d = {}
        s  = d.get("score", 0)
        v  = d.get("verdict", "—")
        c2 = "#16a34a" if s >= 80 else "#ca8a04" if s >= 60 else "#dc2626"
        pills_html += (
            f"<div style='background:{_BG2};border:1px solid {_BORDER};"
            f"border-radius:8px;padding:12px 16px;flex:1;min-width:150px;'>"
            f"<div style='font-size:18px;margin-bottom:4px;'>{icon}</div>"
            f"<div style='font-size:10px;color:{_LBL};font-family:monospace;"
            f"text-transform:uppercase;letter-spacing:1px;'>{label}</div>"
            f"<div style='font-size:24px;font-weight:800;color:{c2};"
            f"font-family:monospace;margin:4px 0;'>{s}</div>"
            f"<div style='font-size:10px;color:{_LBL};font-family:monospace;'>{v}</div>"
            f"</div>"
        )
    st.markdown(
        f"<div style='display:flex;gap:12px;flex-wrap:wrap;margin-bottom:22px;'>"
        f"{pills_html}</div>",
        unsafe_allow_html=True)

    # ── Dimension detail cards ────────────────────────────────────────────────
    st.markdown(_section_header("Dimension Details"), unsafe_allow_html=True)

    def _safe_dim(key: str) -> dict:
        """Return the dimension dict, guarding against str values from malformed LLM output."""
        val = existing.get(key, {})
        return val if isinstance(val, dict) else {}

    _render_ai_assistant_dimension("🎯", "Extraction Accuracy",
                                   _safe_dim("extraction_accuracy"), "#2563eb")
    _render_ai_assistant_dimension("⚡", "Signal Credibility",
                                   _safe_dim("signal_credibility"), "#7c3aed")
    _render_ai_assistant_dimension("📋", "Coverage Analysis",
                                   _safe_dim("coverage_analysis"), "#059669")

    # ── Recommended actions ───────────────────────────────────────────────────
    actions = overall.get("recommended_actions", [])
    if actions:
        st.markdown(_section_header("Recommended Actions"), unsafe_allow_html=True)
        for i, action in enumerate(actions, 1):
            st.markdown(
                f"<div style='background:{_BG2};border:1px solid {_BORDER};"
                f"border-radius:6px;padding:10px 16px;margin-bottom:6px;"
                f"display:flex;gap:12px;align-items:flex-start;'>"
                f"<span style='font-size:11px;font-weight:700;color:#ffffff;"
                f"background:#0284c7;border-radius:50%;width:20px;height:20px;"
                f"display:flex;align-items:center;justify-content:center;"
                f"flex-shrink:0;font-family:monospace;'>{i}</span>"
                f"<span style='font-size:13px;color:{_TXT};line-height:1.6;'>{action}</span>"
                f"</div>",
                unsafe_allow_html=True)

    st.markdown("<div style='margin-top:16px;'></div>", unsafe_allow_html=True)
    _, clr_col = st.columns([5, 2])
    with clr_col:
        if st.button("🗑 Clear results", key=f"_wa_clr_ai_{uploaded_name}",
                     use_container_width=True):
            st.session_state.pop(_VK, None); st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# MAIN ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def render_word_analysis_panel(
    word_result: dict,
    uploaded_name: str,
    file_path: str,
) -> None:
    """
    Main entry point. Call from app2.py:
        render_word_analysis_panel(word_result, uploaded.name, file_path)

    v8 changes vs v7:
      • AI badge removed from field names in Fields tab
      • Signals auto-run on first tab open (no button click required)
      • Legal Document field extraction now captures Case Complaint Summary
        and Causes of Action from the document body paragraphs
      • AttributeError fix: incorrect_fields items handled as both str and dict
      • _safe_dim() guard prevents str.get() crash for malformed LLM responses
    """
    _fp = (
        uploaded_name
        + "|" + str(len(word_result.get("fields", [])))
        + "|" + str(len(word_result.get("raw_text", "")))
    )
    if st.session_state.get("_wa_panel_fp") != _fp:
        for k in list(st.session_state.keys()):
            if k.startswith(_wa_key(uploaded_name, "")):
                st.session_state.pop(k, None)
        st.session_state["_wa_panel_fp"] = _fp

    clf       = word_result.get("doc_classification", {})
    doc_label = clf.get("doc_type_label", "General Document")
    doc_conf  = float(clf.get("confidence", 0.0))
    meta      = _DOC_TYPE_META.get(doc_label, _DOC_TYPE_META["General Document"])
    raw_text  = word_result.get("raw_text", "")

    # ── Augment fields: direct docx + LLM extraction ─────────────────────────
    existing_fields: list[dict] = list(word_result.get("fields", []))
    direct_fields:   list[dict] = []

    if file_path and os.path.exists(file_path):
        direct_fields = _extract_fields_from_docx(file_path)

    _LLM_FIELDS_KEY = _wa_key(uploaded_name, "llm_fields")
    if _LLM_FIELDS_KEY not in st.session_state:
        with st.spinner("Extracting fields using intelligence engine…"):
            llm_fields = _extract_fields_llm(doc_label, raw_text)
        st.session_state[_LLM_FIELDS_KEY] = llm_fields
    llm_fields = st.session_state[_LLM_FIELDS_KEY]

    fields = _merge_fields(existing_fields, direct_fields) if direct_fields else existing_fields
    fields = _merge_fields(fields, llm_fields) if llm_fields else fields

    extra = st.session_state.get(_wa_key(uploaded_name, "extra_fields"), [])
    exist = {f.get("field_name", "") for f in fields}
    for ef in extra:
        if ef.get("field_name") not in exist:
            fields.append(ef)

    # ── Auto-generate summary once per file load ──────────────────────────────
    _SK = _wa_key(uploaded_name, "summary_override")
    if _SK not in st.session_state:
        with st.spinner("Generating document summary…"):
            summary = _auto_summarise_llm(fields, doc_label, raw_text, uploaded_name)
        st.session_state[_SK] = summary

    # ── Panel header ──────────────────────────────────────────────────────────
    st.markdown(
        f"<div style='background:{meta['bg']};border:1px solid {meta['color']}30;"
        f"border-radius:10px;padding:13px 18px;margin-bottom:14px;'>"
        f"<div style='display:flex;align-items:center;gap:12px;'>"
        f"<span style='font-size:22px;'>{meta['icon']}</span>"
        f"<div>"
        f"<div style='font-size:14px;font-weight:700;color:{meta['color']};"
        f"font-family:monospace;text-transform:uppercase;letter-spacing:1.5px;'>"
        f"{doc_label} Document Analysis</div>"
        f"<div style='font-size:11px;color:{_LBL};margin-top:3px;'>"
        f"📄 {html.escape(uploaded_name)} · {len(fields)} field(s) extracted · "
        f"{len(raw_text.split())} words · Confidence: {_conf_badge(doc_conf)}</div>"
        f"</div></div></div>",
        unsafe_allow_html=True)

    # ── Tabs ──────────────────────────────────────────────────────────────────
    tabs = st.tabs([
        f"🔍 Fields ({len(fields)})",
        "📝 Summary",
        "📄 Raw JSON",
        "🔄 Transformation Journey",
        "⚡ Signals",
        "✅ AI Assistant",
    ])
    with tabs[0]: _render_fields_tab(fields, word_result, uploaded_name, file_path)
    with tabs[1]: _render_summary_tab(word_result, fields, uploaded_name)
    with tabs[2]: _render_raw_json_tab(fields, word_result, uploaded_name)
    with tabs[3]: _render_journey_tab(fields, uploaded_name)
    with tabs[4]: _render_signals_tab(fields, word_result, uploaded_name)
    with tabs[5]: _render_ai_assistant_tab(fields, word_result, uploaded_name)